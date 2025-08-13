from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from typing import List, Optional, Union, Dict
from datetime import datetime
from sqlalchemy.orm import Session
from sqlalchemy import text
from app.schemas.document import DocumentBase, DocumentInfo
from app.services.external.s3_service import upload_file, delete_file_from_s3, generate_presigned_url
from app.services.external.postgres_service import save_document, get_documents, get_document_by_id, delete_document_from_postgres
from app.models.documents import Document
from app.services.external.opensearch_service import index_document_chunks, delete_document_chunks_from_opensearch, DOCUMENT_INDEX_NAME
from app.services.core.document_relation_analyzer import document_relation_analyzer

from app.services.core.document_analyzer import document_analyzer
from app.services.core.text2sql_classifier import text2sql_classifier
from app.services.processors.document_type_updater import DocumentTypeUpdater
from app.services.core.document_summarizer import document_summarizer
from app.routers.user_router import get_current_user, get_current_admin_user
from pydantic import BaseModel
import logging
import re

# 파일 처리 관련 라이브러리들
try:
    import pandas as pd
    import io
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

router = APIRouter()
logger = logging.getLogger(__name__)

# 상수 정의
MAX_FILE_SIZE_MB = 10
TRANSACTION_ISOLATION_LEVEL = "SERIALIZABLE"

class TableUploadResult(BaseModel):
    doc_title: str
    doc_type: str
    uploader_id: int
    version: Optional[str]
    created_at: datetime
    message: str
    analysis: Optional[Dict] = None

class BatchUploadResult(BaseModel):
    total_files: int
    successful_uploads: int
    failed_uploads: int
    results: List[Union[DocumentInfo, TableUploadResult]]
    errors: List[Dict[str, str]]


def _extract_csv_data(file_bytes: bytes) -> tuple[str, list]:
    """CSV 파일에서 데이터 추출"""
    if not PANDAS_AVAILABLE:
        raise ImportError("pandas 라이브러리가 설치되지 않았습니다.")
    df = pd.read_csv(io.BytesIO(file_bytes))
    # 모든 컬럼명을 문자열로 변환
    df.columns = df.columns.astype(str)
    return "", df.to_dict('records')

def _extract_excel_data(file_bytes: bytes) -> tuple[str, list]:
    """Excel 파일에서 데이터 추출"""
    if not PANDAS_AVAILABLE:
        raise ImportError("pandas 라이브러리가 설치되지 않았습니다.")
    
    # 일반적인 단일 헤더 처리
    df = pd.read_excel(io.BytesIO(file_bytes))
    # 모든 컬럼명을 문자열로 변환
    df.columns = df.columns.astype(str)
    return "", df.to_dict('records')

def _extract_text_data(file_bytes: bytes) -> tuple[str, list]:
    """TXT 파일에서 텍스트 추출"""
    text = file_bytes.decode("utf-8", errors="ignore")
    return text, []

def _extract_docx_data(file_bytes: bytes) -> tuple[str, list]:
    """DOCX 파일에서 텍스트 추출"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 라이브러리가 설치되지 않았습니다.")
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))

        def _collapse_spaces(s: str) -> str:
            return " ".join(s.split())

        def _normalize_label(s: str) -> str:
            import re
            # 불릿/구분 기호 제거 후 공백 정규화
            s = s.replace('■', ' ').replace(':', ' ').strip()
            label = _collapse_spaces(s)
            # 한글 라벨의 경우 단어 사이 임의 공백 제거(일반화)
            # 예) "연 락 처" → "연락처", "소    속" → "소속"
            if re.fullmatch(r'[가-힣\s]{2,}', label):
                compact = label.replace(' ', '')
                # 너무 길면 원형 유지(영문 혼합/문장 가능성)
                if len(compact) <= 10:
                    label = compact
            return label

        def _normalize_value(s: str) -> str:
            # 값은 기호/콜론 등 보존, 공백만 접기
            return _collapse_spaces(s.strip())

        parts: list[str] = []

        # 1) 일반 문단 텍스트 수집
        paragraph_texts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        if paragraph_texts:
            parts.append("\n".join(paragraph_texts))

        # 2) 표 내용을 텍스트로 평탄화하여 수집
        table_lines: list[str] = []
        seen_recent: set[str] = set()
        for table in doc.tables:
            for row in table.rows:
                raw_cells = [c.text.strip() for c in row.cells]
                cells = [c for c in raw_cells if c and c.strip()]
                if not cells:
                    continue

                # 같은 문구 반복만 있는 행은 한 번만 기록
                if len(set(_normalize_value(c) for c in cells)) == 1:
                    line = _normalize_value(cells[0])
                    if line and line not in seen_recent:
                        table_lines.append(line)
                        seen_recent = {line}
                    continue

                # 짝수 개 컬럼은 라벨:값 페어링으로 변환
                if len(cells) % 2 == 0 and len(cells) <= 8:
                    pairs = []
                    for i in range(0, len(cells), 2):
                        label = _normalize_label(cells[i])
                        value = _normalize_value(cells[i + 1])
                        if label or value:
                            pairs.append(f"{label}: {value}" if value else label)
                    line = " | ".join([p for p in pairs if p])
                else:
                    # 다열 표는 공백/중복 제거 후 결합
                    non_empty = []
                    for c in cells:
                        val = _normalize_value(c)
                        if val and (not non_empty or val != non_empty[-1]):
                            non_empty.append(val)
                    line = " | ".join(non_empty)

                line = line.strip()
                if line and (not table_lines or line != table_lines[-1]):
                    table_lines.append(line)

            # 표 간 구분 빈 줄 추가
            if table_lines and table_lines[-1] != "":
                table_lines.append("")

        if table_lines:
            parts.append("\n".join(table_lines))

        text = "\n".join(parts).strip()
        return text, []
    except Exception as e:
        logger.error(f"DOCX 파일 텍스트 추출 실패: {e}")
        return "", []

def _extract_pdf_data(file_bytes: bytes) -> tuple[str, list]:
    """PDF 파일에서 텍스트 추출"""
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 라이브러리가 설치되지 않았습니다.")
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text, []
    except Exception as e:
        logger.error(f"PDF 파일 텍스트 추출 실패: {e}")
        return "", []

# 파일 처리 관련 상수 (함수 정의 후에 배치)
FILE_PROCESSORS = {
    '.csv': _extract_csv_data,
    '.xlsx': _extract_excel_data,
    '.xls': _extract_excel_data,
    '.txt': _extract_text_data,
    '.docx': _extract_docx_data,
    '.pdf': _extract_pdf_data,
}

def validate_file_size(file: UploadFile, max_size_mb: int = MAX_FILE_SIZE_MB) -> bytes:
    """
    파일 크기를 검증하고 파일 바이트를 반환합니다.
    
    Args:
        file: 업로드된 파일
        max_size_mb: 최대 파일 크기 (MB)
        
    Returns:
        bytes: 파일 바이트 데이터
        
    Raises:
        HTTPException: 파일 크기 초과 시
    """
    file.file.seek(0, 2)
    file_size = file.file.tell()
    file.file.seek(0)
    
    if file_size > max_size_mb * 1024 * 1024:
        raise HTTPException(
            status_code=400, 
            detail=f"파일 크기가 너무 큽니다. 최대 {max_size_mb}MB까지 업로드 가능합니다."
        )
    
    return file.file.read()

def extract_doc_title(filename: str) -> str:
    """
    파일명에서 문서 제목을 추출합니다.
    
    Args:
        filename: 파일명
        
    Returns:
        str: 문서 제목 (확장자 제외)
    """
    return filename.rsplit('.', 1)[0] if '.' in filename else filename

def extract_text_and_table(file_bytes: bytes, filename: str):
    """
    파일 확장자에 따라 텍스트/테이블 데이터를 추출한다.
    
    Args:
        file_bytes: 파일 바이트 데이터
        filename: 파일명
        
    Returns:
        tuple: (text, table_data, is_table_file)
        
    Raises:
        HTTPException: 지원하지 않는 파일 형식 또는 처리 오류
    """
    file_extension = document_analyzer._get_file_extension(filename)
    is_table_file = file_extension in document_analyzer.supported_extensions["table"]
    
    text = ""
    table_data = []
    
    try:
        processor = FILE_PROCESSORS.get(file_extension.lower())
        if processor:
            text, table_data = processor(file_bytes)
        else:
            raise HTTPException(status_code=400, detail=f"지원하지 않는 파일 형식입니다: {file_extension}")
            
    except ImportError as e:
        logger.warning(f"필요한 라이브러리가 설치되지 않았습니다: {e}")
        raise HTTPException(status_code=500, detail="파일 처리를 위한 라이브러리가 필요합니다.")
    except Exception as e:
        logger.error(f"파일 텍스트 추출 실패: {e}")
        raise HTTPException(status_code=500, detail=f"파일 처리 중 오류가 발생했습니다: {str(e)}")
    
    return text, table_data, is_table_file

async def process_table_document(
    file_bytes: bytes,
    filename: str,
    doc_title: str,
    table_data: list,
    uploader_id: int,
    version: str = None,
    session: Session = None
) -> TableUploadResult:
    """
    테이블 문서를 처리합니다.
    
    Args:
        file_bytes: 파일 바이트 데이터
        filename: 파일명
        doc_title: 문서 제목
        table_data: 테이블 데이터
        uploader_id: 업로더 ID
        version: 문서 버전
        session: 데이터베이스 세션
        
    Returns:
        TableUploadResult: 테이블 업로드 결과
    """
    logger.info(f"테이블 문서 Text2SQL 처리 시작: {filename}")
    
    # Text2SQL 분류기로 처리
    result = await text2sql_classifier.classify_table_with_text2sql(
        table_data=table_data,
        table_description=doc_title,
        document_id=None,
        uploader_id=uploader_id
    )
    
    if not result['success']:
        logger.error(f"Text2SQL 분류 실패: {result['message']}")
        raise HTTPException(
            status_code=500,
            detail=f"문서 분류 중 오류가 발생했습니다: {result['message']}"
        )
    
    logger.info(f"Text2SQL 분류 완료: {result['message']}")
    logger.info(f"분류 결과: {filename} -> {result['target_table']} (신뢰도: {result['confidence']:.2f})")
    
    # 문서 요약 생성 (비동기 처리, 실패해도 계속 진행)
    summary = None
    try:
        doc_type = f"text2sql_{result['target_table']}"
        summary = document_summarizer.summarize_table_document(table_data, doc_title, doc_type)
        if summary:
            logger.info(f"테이블 문서 요약 생성 성공: {doc_title}")
        else:
            logger.warning(f"테이블 문서 요약 생성 실패 (빈 응답): {doc_title}")
    except Exception as e:
        logger.error(f"테이블 문서 요약 생성 중 오류 (계속 진행): {e}")
        summary = None
    
    # S3에 파일 저장
    file_path = upload_file(file_bytes, filename, "application/octet-stream")
    
    # 문서 메타데이터 생성 (요약 포함)
    meta = DocumentBase(
        doc_title=doc_title,
        doc_type=f"text2sql_{result['target_table']}",
        file_path=file_path,
        uploader_id=uploader_id,
        version=version,
        summary=summary,
        created_at=datetime.now()
    )
    
    # 세션이 제공된 경우 세션 사용, 아니면 새로 저장
    if session:
        db_doc = Document(**meta.dict())
        session.add(db_doc)
        session.flush()
    else:
        db_doc = save_document(meta)
    
    # 문서 타입 업데이트
    try:
        if session:
            await DocumentTypeUpdater.update_after_success(db_doc, result, session)
        else:
            from app.services.utils.db import create_db_session
            with create_db_session() as update_session:
                doc_to_update = get_document_by_id(db_doc.doc_id)
                if doc_to_update:
                    await DocumentTypeUpdater.update_after_success(doc_to_update, result, update_session)
                    update_session.commit()
        logger.info(f"문서 타입 업데이트 완료: {db_doc.doc_id}")
    except Exception as e:
        logger.error(f"문서 타입 업데이트 실패: {e}")
    
    logger.info(f"테이블 문서 업로드 완료: {db_doc.doc_id}")
    
    return TableUploadResult(
        doc_title=doc_title,
        doc_type=f"text2sql_{result['target_table']}",
        uploader_id=uploader_id,
        version=version,
        created_at=datetime.now(),
        message=f"{result['message']} (문서 ID: {db_doc.doc_id})",
        analysis={
            'target_table': result['target_table'],
            'confidence': result['confidence'],
            'reasoning': result.get('reasoning', ''),
            'column_mapping': result.get('column_mapping', {}),
            'doc_id': db_doc.doc_id
        }
    )

async def process_text_document(
    file_bytes: bytes,
    filename: str,
    doc_title: str,
    text: str,
    file_extension: str,
    uploader_id: int,
    version: str = None,
    session: Session = None
) -> DocumentInfo:
    """
    텍스트 문서를 처리합니다.
    
    Args:
        file_bytes: 파일 바이트 데이터
        filename: 파일명
        doc_title: 문서 제목
        text: 추출된 텍스트
        file_extension: 파일 확장자
        uploader_id: 업로더 ID
        version: 문서 버전
        session: 데이터베이스 세션
        
    Returns:
        DocumentInfo: 문서 정보
    """
    logger.info(f"텍스트 문서 처리 시작: {filename}")
    
    # 문서 타입 분석
    analyzed_doc_type = document_analyzer.analyze_document(text, filename)
    logger.info(f"문서 분석 결과: {filename} -> {analyzed_doc_type}")
    
    # 문서 요약 생성 (비동기 처리, 실패해도 계속 진행)
    summary = None
    try:
        summary = document_summarizer.summarize_text_document(text, doc_title, analyzed_doc_type)
        if summary:
            logger.info(f"문서 요약 생성 성공: {doc_title}")
        else:
            logger.warning(f"문서 요약 생성 실패 (빈 응답): {doc_title}")
    except Exception as e:
        logger.error(f"문서 요약 생성 중 오류 (계속 진행): {e}")
        summary = None
    
    # S3 업로드
    file_path = upload_file(file_bytes, filename, "application/octet-stream")
    
    # 문서 메타데이터 생성 (요약 포함)
    meta = DocumentBase(
        doc_title=doc_title,
        doc_type=analyzed_doc_type,
        file_path=file_path,
        uploader_id=uploader_id,
        version=version,
        summary=summary,
        created_at=datetime.now()
    )
    
    # 세션이 제공된 경우 세션 사용, 아니면 새로 저장
    if session:
        db_doc = Document(**meta.dict())
        session.add(db_doc)
        session.flush()
    else:
        db_doc = save_document(meta)
    
    # 문서 타입에 따른 처리 분기
    if file_extension in document_analyzer.supported_extensions["text"]:
        # 규정/법률 문서는 OpenSearch에 청킹하여 저장
        if analyzed_doc_type in ["regulation", "law"]:
            chunking_type = document_analyzer.get_chunking_type(analyzed_doc_type)
            index_document_chunks(
                doc_id=db_doc.doc_id,
                doc_title=doc_title,
                file_name=filename,
                text=text,
                document_type=chunking_type
            )
            logger.info(f"규정/법률 문서 업로드 완료: {db_doc.doc_id} (타입: {analyzed_doc_type}, 청킹: {chunking_type})")
        
        # 보고서 문서는 관계 분석 후 DocumentRelation에 저장
        elif analyzed_doc_type == "report":
            # 문서 관계 분석
            relation_result = document_relation_analyzer.analyze_document_relations(
                doc_id=db_doc.doc_id,
                text=text,
                table_data=None
            )
            
            if relation_result['success']:
                logger.info(f"보고서 문서 업로드 완료: {db_doc.doc_id} (타입: {analyzed_doc_type}, 관계: {relation_result['relations_created']}개)")
            else:
                logger.warning(f"보고서 문서 관계 분석 실패: {relation_result['message']}")
                logger.info(f"보고서 문서 업로드 완료: {db_doc.doc_id} (타입: {analyzed_doc_type})")
        else:
            logger.info(f"문서 업로드 완료: {db_doc.doc_id} (타입: {analyzed_doc_type})")
    else:
        logger.info(f"문서 업로드 완료: {db_doc.doc_id} (타입: {analyzed_doc_type})")
    
    return DocumentInfo.model_validate(db_doc)

async def process_single_document(file: UploadFile, uploader_id: int, version: str = None) -> Union[DocumentInfo, TableUploadResult]:
    """
    단일 문서를 처리하는 공통 함수
    
    Args:
        file: 업로드할 파일
        uploader_id: 업로더 ID
        version: 문서 버전 (선택사항)
        
    Returns:
        DocumentInfo 또는 TableUploadResult: 처리 결과
    """
    # 세션 없이 호출 (기본 동작)
    return await process_single_document_with_session(file, uploader_id, version, None)


@router.post("/documents/upload", response_model=Union[DocumentInfo, TableUploadResult])
async def upload_document(file: UploadFile = File(...), doc_title: str = Form(None), uploader_id: int = Form(...), version: str = Form(None), user=Depends(get_current_user)):
    """
    문서를 업로드하고 자동으로 타입을 분석하여 저장합니다.
    
    Args:
        file: 업로드할 파일
        doc_title: 문서 제목 (선택사항, 없으면 파일명 사용)
        uploader_id: 업로더 ID
        version: 문서 버전 (선택사항)
        user: 현재 인증된 사용자
        
    Returns:
        DocumentInfo 또는 TableUploadResult: 업로드 결과
        
    Raises:
        HTTPException: 파일 크기 초과, 지원하지 않는 형식, 처리 오류 등
    """
    try:
        return await process_single_document(file, uploader_id, version)
    except Exception as e:
        logger.error(f"문서 업로드 실패: {e}")
        raise HTTPException(status_code=500, detail=f"문서 업로드 중 오류가 발생했습니다: {str(e)}")

@router.post("/documents/upload/batch", response_model=BatchUploadResult)
async def upload_documents_batch(files: List[UploadFile] = File(...), uploader_id: int = Form(...), version: str = Form(None), user=Depends(get_current_user)):
    """
    여러 문서를 한 번에 업로드합니다.
    
    Args:
        files: 업로드할 파일들
        uploader_id: 업로더 ID
        version: 문서 버전 (선택사항)
        user: 현재 인증된 사용자
        
    Returns:
        BatchUploadResult: 배치 업로드 결과
    """
    total_files = len(files)
    successful_uploads = 0
    failed_uploads = 0
    results = []
    errors = []
    
    # 배치 업로드 시작 로그
    logger.info(f"배치 업로드 시작: {total_files}개 파일")
    
    # 각 파일을 개별 트랜잭션으로 처리하여 격리 보장
    for i, file in enumerate(files, 1):
        try:
            logger.info(f"배치 업로드 진행 중: {i}/{total_files} - {file.filename}")
            
            # 개별 파일 처리를 위한 트랜잭션 격리
            from app.services.utils.db import create_db_session
            with create_db_session() as session:
                try:
                    # 트랜잭션 격리 레벨 설정
                    session.execute(text(f"SET TRANSACTION ISOLATION LEVEL {TRANSACTION_ISOLATION_LEVEL}"))
                    
                    # 파일 처리 (세션 전달)
                    result = await process_single_document_with_session(file, uploader_id, version, session)
                    results.append(result)
                    successful_uploads += 1
                    logger.info(f"배치 업로드 성공: {file.filename}")
                    
                    session.commit()
                    
                except Exception as e:
                    session.rollback()
                    raise e
                    
        except HTTPException as e:
            # HTTP 예외는 상세 메시지 보존
            errors.append({"filename": file.filename, "error": e.detail})
            failed_uploads += 1
            logger.error(f"배치 업로드 중 오류 ({file.filename}): {e.detail}")
        except Exception as e:
            # 기타 예외
            error_msg = f"문서 업로드 실패: {str(e)}"
            errors.append({"filename": file.filename, "error": error_msg})
            failed_uploads += 1
            logger.error(f"배치 업로드 중 오류 ({file.filename}): {e}")
        
        # 개별 파일 실패 시에도 계속 진행
    
    logger.info(f"배치 업로드 완료: 성공 {successful_uploads}/{total_files}, 실패 {failed_uploads}")
    
    return BatchUploadResult(
        total_files=total_files,
        successful_uploads=successful_uploads,
        failed_uploads=failed_uploads,
        results=results,
        errors=errors
    )

async def process_single_document_with_session(file: UploadFile, uploader_id: int, version: str = None, session: Session = None) -> Union[DocumentInfo, TableUploadResult]:
    """
    세션을 받아서 단일 문서를 처리하는 함수 (트랜잭션 격리용)
    
    Args:
        file: 업로드할 파일
        uploader_id: 업로더 ID
        version: 문서 버전
        session: 데이터베이스 세션 (옵션)
        
    Returns:
        DocumentInfo 또는 TableUploadResult: 처리 결과
    """
    # 파일 크기 검증
    file_bytes = validate_file_size(file)
    
    # 파일 분석
    file_extension = document_analyzer._get_file_extension(file.filename)
    text, table_data, is_table_file = extract_text_and_table(file_bytes, file.filename)
    
    # 문서 제목 추출
    doc_title = extract_doc_title(file.filename)
    
    # 테이블 문서 처리
    if is_table_file and table_data:
        return await process_table_document(
            file_bytes=file_bytes,
            filename=file.filename,
            doc_title=doc_title,
            table_data=table_data,
            uploader_id=uploader_id,
            version=version,
            session=session
        )
    
    # 텍스트 문서 처리
    else:
        return await process_text_document(
            file_bytes=file_bytes,
            filename=file.filename,
            doc_title=doc_title,
            text=text,
            file_extension=file_extension,
            uploader_id=uploader_id,
            version=version,
            session=session
        )

@router.get("/documents/", response_model=List[DocumentInfo])
def list_documents(user=Depends(get_current_user)):
    """
    모든 문서 목록을 조회합니다.
    
    Args:
        user: 현재 인증된 사용자
        
    Returns:
        List[DocumentInfo]: 문서 목록
    """
    docs = get_documents()
    return [DocumentInfo.model_validate(doc) for doc in docs]

@router.get("/documents/{doc_id}")
def get_document(doc_id: int, user=Depends(get_current_user)):
    """
    특정 문서를 조회합니다. 다운로드 링크를 포함합니다.
    
    Args:
        doc_id: 문서 ID
        user: 현재 인증된 사용자
        
    Returns:
        Dict: 문서 정보 및 다운로드 링크
        
    Raises:
        HTTPException: 문서를 찾을 수 없는 경우
    """
    doc = get_document_by_id(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # 문서 정보를 dict로 변환
    doc_info = DocumentInfo.model_validate(doc).dict()
    
    # S3에서 파일명 추출 (URL에서 파일명 부분만 추출)
    file_name = doc.file_path.split("/")[-1]
    
    # Pre-signed URL 생성 (1시간 유효)
    download_url = generate_presigned_url(file_name, expiration=3600)
    
    # 다운로드 링크 추가
    doc_info["download_url"] = download_url
    doc_info["download_expires_in"] = "1 hour"
    
    logger.info(f"문서 상세 조회 완료 (다운로드 링크 포함): doc_id={doc_id}")
    
    return doc_info

@router.get("/documents/{doc_id}/download")
def get_document_download_link(doc_id: int, expiration_hours: int = 1, user=Depends(get_current_user)):
    """
    문서의 다운로드 링크를 생성합니다.
    
    Args:
        doc_id: 문서 ID
        expiration_hours: 링크 유효 시간 (시간 단위, 기본값: 1시간, 최대: 24시간)
        user: 현재 인증된 사용자
        
    Returns:
        Dict: 다운로드 링크 정보
        
    Raises:
        HTTPException: 문서를 찾을 수 없는 경우
    """
    # 유효 시간 제한 (최대 24시간)
    if expiration_hours > 24:
        expiration_hours = 24
    elif expiration_hours < 1:
        expiration_hours = 1
    
    doc = get_document_by_id(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # S3에서 파일명 추출
    file_name = doc.file_path.split("/")[-1]
    
    # Pre-signed URL 생성
    expiration_seconds = expiration_hours * 3600
    download_url = generate_presigned_url(file_name, expiration=expiration_seconds)
    
    if not download_url:
        raise HTTPException(status_code=500, detail="Failed to generate download link")
    
    logger.info(f"다운로드 링크 생성 완료: doc_id={doc_id}, 유효시간={expiration_hours}시간")
    
    return {
        "doc_id": doc_id,
        "doc_title": doc.doc_title,
        "file_name": file_name,
        "download_url": download_url,
        "expires_in_hours": expiration_hours,
        "generated_at": datetime.now().isoformat()
    }

@router.delete("/documents/{doc_id}", response_model=DocumentInfo)
def delete_document(doc_id: int, admin=Depends(get_current_admin_user)):
    """
    문서를 삭제합니다. (관리자만 가능)
    
    Args:
        doc_id: 삭제할 문서 ID
        admin: 현재 인증된 관리자
        
    Returns:
        DocumentInfo: 삭제된 문서 정보
        
    Raises:
        HTTPException: 문서를 찾을 수 없거나 삭제 실패 시
    """
    doc = get_document_by_id(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # 1. S3에서 원본 파일 삭제
        file_name = doc.file_path.split("/")[-1]
        delete_file_from_s3(file_name)
        logger.info(f"S3 파일 삭제 완료: {file_name}")
        
        # 2. OpenSearch에서 문서 청크 삭제 (텍스트 문서용)
        delete_document_chunks_from_opensearch(DOCUMENT_INDEX_NAME, doc_id)
        logger.info(f"OpenSearch 문서 청크 삭제 완료: doc_id={doc_id}")
        

        
        # 4. 문서 관계 삭제
        try:
            relation_delete_result = document_relation_analyzer.delete_document_relations(doc_id)
            if relation_delete_result['success']:
                logger.info(f"문서 관계 삭제 완료: {relation_delete_result['deleted_count']}개")
            else:
                logger.warning(f"문서 관계 삭제 실패: {relation_delete_result['message']}")
        except Exception as e:
            logger.error(f"문서 관계 삭제 중 오류: {e}")
        
        # 5. PostgreSQL에서 문서 메타데이터 삭제
        deleted_doc = delete_document_from_postgres(doc_id)
        if not deleted_doc:
            raise HTTPException(status_code=500, detail="Failed to delete document from DB")
        
        logger.info(f"문서 완전 삭제 완료: doc_id={doc_id}, title={doc.doc_title}")
        return DocumentInfo.model_validate(deleted_doc)
        
    except Exception as e:
        logger.error(f"문서 삭제 중 오류: {e}")
        raise HTTPException(status_code=500, detail=f"문서 삭제 중 오류가 발생했습니다: {str(e)}") 