from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage
from langchain_core.prompts import ChatPromptTemplate
import json
import yaml
import time
import re
import os
import sys
import uuid
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
import requests
from typing import Dict, List, Tuple, Optional, Any

# 동적으로 backend 디렉토리 찾기
current_file = Path(__file__).resolve()
backend_dir = None

# backend 디렉토리를 찾을 때까지 상위로 이동
for parent in current_file.parents:
    if parent.name == 'backend' and (parent / 'app').exists():
        backend_dir = parent
        break

if backend_dir:
    sys.path.insert(0, str(backend_dir))

# 이제 import 시도
try:
    from app.core.config import config
    from app.services.tools.common_tools import check_phrase_against_regulations
except ImportError as e:
    print(f"Import Error: {e}")
    print(f"Current working directory: {os.getcwd()}")
    print(f"Python path: {sys.path}")
    # config가 없을 경우 기본값 사용
    class config:
        @staticmethod
        def get_database_api_url():
            # Docker 환경에서는 컨테이너명으로 통신
            # AWS 환경에서는 환경변수로 설정된 URL 사용
            database_api_url = os.getenv('DATABASE_API_URL')
            
            if database_api_url:
                # 환경변수가 설정된 경우 (AWS 또는 프로덕션)
                return database_api_url
            elif os.path.exists('/.dockerenv'):
                # Docker 컨테이너 내부에서 실행 중인 경우
                return 'http://database_api:8000'
            else:
                # 로컬 개발 환경
                return 'http://localhost:8000'

load_dotenv()

class CreateDocumentFunction:
    def __init__(self, model_name: str = "gpt-4o-mini", temperature: float = 0.7):
        self.model_name = model_name
        self.temperature = temperature

        # LLM 초기화
        self.llm = ChatOpenAI(
            model=self.model_name, 
            temperature=self.temperature
        )

        # YAML 파일에서 템플릿 로드
        self.doc_prompts = self._load_templates()

    def _load_templates(self):
        """
        YAML 파일에서 문서 템플릿 및 프롬프트 정보를 로드합니다.
        
        템플릿 파일 구조:
        - templates.yaml 파일에서 각 문서 타입별 템플릿 정보 로드
        - 각 문서 타입마다 input_prompt, choan_system_prompt, choan_fallback_fields 포함
        
        Returns:
            dict: 문서 타입별 템플릿 정보 딕셔너리
                 예: {"영업방문 결과보고서": {"input_prompt": "...", "choan_system_prompt": "...", ...}}
                 로드 실패 시 빈 딕셔너리 반환
        
        Raises:
            Exception: 파일 읽기 실패 시 예외 처리하고 빈 딕셔너리 반환
        """
        try:
            # 현재 스크립트와 같은 디렉토리에서 templates.yaml 파일 찾기
            current_dir = Path(__file__).parent
            template_path = current_dir / "templates.yaml"
            
            # 템플릿 파일 존재 여부 확인
            if not template_path.exists():
                print(f"[WARNING] 템플릿 파일을 찾을 수 없습니다: {template_path}")
                return {}
            
            # YAML 파일 읽기 및 파싱
            with open(template_path, 'r', encoding='utf-8') as file:
                data = yaml.safe_load(file)
                return data.get('templates', {})
                
        except Exception as e:
            print(f"[ERROR] 템플릿 로드 중 오류 발생: {e}")
            return {}

    def input_data(self, request_data: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """
        프론트엔드에서 받은 데이터를 처리하여 문서를 생성합니다.
        
        Args:
            request_data (dict): 프론트엔드에서 받은 요청 데이터
                - type (str): 문서 타입
                - content (dict): 문서 내용 데이터
        """

        document_type = request_data.get('type', '')
        content_data = request_data.get('content', {})
        
        return document_type, content_data
    
    def policy_violation_content_extraction(self, document_type: str, content_data: Dict[str, Any]) -> List[str]:
        """
        문서 타입에 따라 정책 위반 검토가 필요한 내용을 추출합니다.
        
        Args:
            document_type (str): 문서 타입 ('영업방문결과보고서', '제품설명회시행계획서', '제품설명회결과보고서')
            content_data (dict): 문서 내용 데이터
            
        Returns:
            list: 추출된 내용 리스트
        """
        result = []
        
        if document_type == '영업방문결과보고서':
            # 영업방문결과보고서의 경우 특정 키만 처리
            target_keys = ['고객사개요', '프로젝트개요', '방문및협의내용', '향후계획및일정', '협조사항및공유사항']
            for key, value in content_data.items():
                if key in target_keys:
                    result.append(f'{key}는 {value}')
                    
        elif document_type == '제품설명회시행신청서':
            # 제품설명회시행계획서의 경우 특정 키만 처리
            target_keys = ['제품설명회시행목적', '제품설명회주요내용']
            for key, value in content_data.items():
                if key in target_keys:
                    result.append(f'{key}는 {value}')
                    
        elif document_type == '제품설명회시행결과보고서':
            # 제품설명회결과보고서의 경우
            general_keys = ['제품설명회시행목적', '제품설명회주요내용', '지급내역']
            budget_keys = ['금액', '메뉴', '주류', '1인금액']
            
            for key, value in content_data.items():
                if key in general_keys:
                    result.append(f'{key}는 {value}')
                elif key in budget_keys:
                    result.append(f'예산 사용 내역 {key}는 {value}')
                    
        return result

    def check_policy_violation(self, content_list: List[str]) -> str:
        violations = []
        fastapi_url = f"{config.get_database_api_url()}/qa/question"
        
        for phrase in content_list:
            try:
                # FastAPI 호출 - 올바른 페이로드 형식 사용
                payload = {
                    "question": phrase,
                    "top_k": 5,
                    "include_summary": True,
                    "include_sources": True
                }
                
                response = requests.post(
                    fastapi_url,
                    json=payload,
                    headers={"Content-Type": "application/json"},
                    timeout=30
                )
                
                if response.status_code == 200:
                    api_result = response.json()
                    if api_result.get('success', False):
                        search_results = api_result.get('search_results', [])
                        print(f"📊 '{phrase}' 검색 결과: {len(search_results)}개")
                        
                        # 3단계: LLM을 사용해 추출된 규정 정보와 비교하여 위반 여부 판단
                        violation_result = check_phrase_against_regulations(phrase, search_results, self.llm)
                        if violation_result != "OK":
                            violations.append(violation_result)
                    else:
                        print(f"⚠️ API 응답 실패 ({phrase}): {api_result}")
                        violations.append(f"{phrase}: API 응답 오류")
                        
                else:
                    print(f"⚠️ FastAPI 호출 실패 ({phrase}): {response.status_code}")
                    violations.append(f"{phrase}: 규정 검색 실패 (HTTP {response.status_code})")
                    
            except requests.exceptions.RequestException as e:
                print(f"⚠️ API 호출 오류 ({phrase}): {e}")
                violations.append(f"{phrase}: 네트워크 오류로 규정 확인 불가")
            except Exception as e:
                print(f"⚠️ 처리 중 오류 ({phrase}): {e}")
                violations.append(f"{phrase}: 처리 오류 - {str(e)}")
        
        # 최종 결과 반환
        actual_violations = []
        for violation in violations:
            # "OK"가 포함되지 않은 실제 위반 사항만 추가
            if ": OK" not in violation and violation.strip() != "OK":
                actual_violations.append(violation)
        
        if actual_violations:
            return "\n\n".join(actual_violations)
        else:
            return "OK"
        
    def convert_structured_to_natural_text(self, document_type: str, content_data: Dict[str, Any]) -> str:
        """
        문서 타입에 따라 구조화된 데이터를 자연어 텍스트로 변환합니다.
        
        Args:
            document_type (str): 문서 타입
            content_data (Dict[str, Any]): 구조화된 데이터
            
        Returns:
            str: 자연어로 변환된 텍스트
        """
        if document_type == "영업방문결과보고서":
            # 영업방문결과보고서의 경우: f'{key}는 {value}' 형태로 변환
            text_parts = []
            for key, value in content_data.items():
                if value:  # 값이 있는 경우만 포함
                    text_parts.append(f'{key}는 {value}')
            
            # 리스트를 ', '로 연결하여 하나의 텍스트로 반환
            return ', '.join(text_parts)
        
        elif document_type in ["제품설명회시행신청서", "제품설명회시행결과보고서"]:
            # 제품설명회 관련 문서의 경우
            text_parts = []
            for key, value in content_data.items():
                if value:  # 값이 있는 경우만 포함
                    # 리스트 형태의 데이터 처리
                    if isinstance(value, list):
                        # 이중 리스트인 경우 (예: [['영업팀', '손현성'], ['영업팀', '허한결']])
                        if all(isinstance(item, list) for item in value):
                            converted_items = []
                            for item in value:
                                if len(item) >= 2:
                                    # '팀명 이름' 형식으로 변환
                                    converted_items.append(f'{item[0]} {item[1]}')
                                else:
                                    # 리스트가 2개 미만의 요소를 가진 경우
                                    converted_items.append(' '.join(str(x) for x in item))
                            value_str = ', '.join(converted_items)
                        else:
                            # 단일 리스트인 경우
                            value_str = ', '.join(str(item) for item in value)
                    else:
                        # 일반 문자열 또는 숫자인 경우
                        value_str = str(value)
                    
                    text_parts.append(f'{key}는 {value_str}')
            
            # 리스트를 ', '로 연결하여 하나의 텍스트로 반환
            return ', '.join(text_parts)
        
        # 기본값: JSON 문자열로 반환
        return json.dumps(content_data, ensure_ascii=False)
    
    def parse_user_input(self, document_type: str, content_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        LLM을 사용하여 사용자 입력 데이터를 구조화된 형식으로 파싱합니다.
        
        Args:
            document_type (str): 문서 타입
            content_data (Dict[str, Any]): 사용자가 입력한 원본 데이터
            
        Returns:
            Dict[str, Any]: LLM이 파싱한 구조화된 데이터
        """

        # 내부 메서드로 구조화된 데이터를 자연어로 변환
        natural_text = self.convert_structured_to_natural_text(document_type, content_data)
        print(f"[INFO] 구조화된 데이터를 자연어로 변환 성공")
        print(f"[INFO] 변환된 텍스트: {natural_text[:200]}...")
        user_input = natural_text
            
        # 프롬프트 확인 - 템플릿 파일의 키 이름과 매핑
        template_key_mapping = {
            "영업방문결과보고서": "영업방문 결과보고서",
            "제품설명회시행신청서": "제품설명회 시행 신청서",
            "제품설명회시행결과보고서": "제품설명회 시행 결과보고서"
        }
        
        template_key = template_key_mapping.get(document_type, document_type)
        
        if template_key not in self.doc_prompts:
            print(f"[WARNING] {template_key}에 대한 프롬프트가 없습니다. 원본 데이터 반환")
            return content_data
            
        system_prompt = self.doc_prompts[template_key].get("choan_system_prompt", "")
        if not system_prompt:
            print(f"[WARNING] {document_type}에 대한 시스템 프롬프트가 없습니다. 원본 데이터 반환")
            return content_data

        # 중괄호 이스케이프 처리
        escaped_input = user_input.replace("{", "{{").replace("}", "}}")
        
        parsing_prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{user_input}")
        ])

        try:
            formatted_messages = parsing_prompt.format_messages(user_input=escaped_input)
            print("[INFO] LLM에 전달된 메시지:")
            for m in formatted_messages:
                print(f"[{m.type.upper()}] {m.content[:200]}...")

            response = self.llm.invoke(formatted_messages)
            content = response.content
            json_str = content if isinstance(content, str) else str(content)
            
            print(f"\n[INFO] LLM 응답 내용:\n{json_str[:500]}...")

            # JSON 추출 및 파싱
            if "{" in json_str and "}" in json_str:
                start = json_str.find("{")
                end = json_str.rfind("}") + 1
                clean_json = json_str[start:end]
                
                print(f"\n[INFO] 추출된 JSON:\n{clean_json}")

                try:
                    parsed_data = json.loads(clean_json)
                    print("[SUCCESS] 파싱 성공:", parsed_data)
                    return parsed_data
                    
                except json.JSONDecodeError as json_error:
                    print(f"[ERROR] JSON 파싱 오류: {json_error}")
                    print(f"파싱 시도한 JSON: {repr(clean_json)}")
                    # 파싱 실패 시 원본 데이터 반환
                    return content_data
            else:
                print("[WARNING] 구조화된 JSON 형식을 찾을 수 없음. 원본 데이터 반환")
                return content_data

        except Exception as e:
            print(f"\n[ERROR] 예외 발생: {e}")
            print("[WARNING] LLM 파싱 실패. 원본 데이터 반환")
            return content_data

        
    def create_document(self, document_type: str, content_data: Dict[str, Any]) -> Optional[str]:
        # 문서 타입에 따른 템플릿 파일 매핑
        # 프론트엔드에서 보내는 타입명과 매칭
        template_mapping = {
            "영업방문 결과보고서": "영업방문 결과보고서(템플릿형).docx",
            "영업방문결과보고서": "영업방문 결과보고서(템플릿형).docx",  # 공백 없는 버전도 지원
            "제품설명회 시행 신청서": "제품설명회 시행 신청서(템플릿형).docx",
            "제품설명회시행신청서": "제품설명회 시행 신청서(템플릿형).docx",  # 공백 없는 버전도 지원
            "제품설명회 시행 결과보고서": "제품설명회 시행 결과보고서(템플릿형).docx",
            "제품설명회시행결과보고서": "제품설명회 시행 결과보고서(템플릿형).docx"  # 공백 없는 버전도 지원
        }
        template_filename = template_mapping.get(document_type)
        if not template_filename:
            print(f"ERROR: 지원하지 않는 문서 타입: {document_type}")
            return None
        
        # S3 폴더에서 템플릿 파일 경로 구성
        current_dir = Path(__file__).parent
        template_path = current_dir / "S3" / template_filename
        
        if not template_path.exists():
            print(f"ERROR: 템플릿 파일을 찾을 수 없습니다: {template_path}")
            return None
        
        try:
            # 템플릿 파일 읽기
            print(f"LOADING: 템플릿 파일 로딩: {template_filename}")
            doc = Document(str(template_path))
            
            print(f"PROCESSING: 템플릿 플레이스홀더 치환 중...")
            
            # 양식을 유지하면서 플레이스홀더만 치환
            self._replace_placeholders_in_document(doc, content_data, document_type)
            
            # agent_result_folder 디렉토리 생성
            result_folder = current_dir / "agent_result_folder"
            result_folder.mkdir(exist_ok=True)
            
            # 완성된 문서 저장
            today_date = time.strftime('%Y%m%d')
            timestamp = time.strftime('%H%M%S')
            doc_type_no_space = document_type.replace(" ", "")
            output_filename = f"{doc_type_no_space}_{today_date}_{timestamp}.docx"
            output_path = result_folder / output_filename
            
            # 파일이 이미 존재하는 경우 처리
            if output_path.exists():
                # 고유한 파일명 생성 (밀리초 추가)
                import datetime
                milliseconds = datetime.datetime.now().strftime('%f')[:3]
                output_filename = f"{doc_type_no_space}_{today_date}_{timestamp}_{milliseconds}.docx"
                output_path = result_folder / output_filename
            
            try:
                doc.save(str(output_path))
            except PermissionError:
                # 권한 오류 발생 시 임시 파일명으로 재시도
                import tempfile
                temp_filename = f"{doc_type_no_space}_{today_date}_{timestamp}_temp_{uuid.uuid4().hex[:8]}.docx"
                output_path = result_folder / temp_filename
                doc.save(str(output_path))
            
            print("SUCCESS: 문서 생성 및 저장 완료!")
            print(f"SAVED_PATH: 저장 경로: {output_path}")
            print("INFO: 템플릿 양식이 그대로 유지되면서 플레이스홀더만 치환되었습니다.")
            
            return str(output_path)
            
        except Exception as e:
            print(f"ERROR: 문서 생성 중 오류 발생: {e}")
            return None

    def _replace_placeholders_in_document(self, doc, filled_data, doc_type):
        """
        문서의 플레이스홀더를 실제 데이터로 치환합니다.
        
        Args:
            doc: DOCX 문서 객체
            filled_data (dict): 치환할 데이터
            doc_type (str): 문서 타입
        
        Returns:
            None: 문서 객체를 직접 수정
        """
        
        # 다중 항목 처리를 위한 특별 처리 필요 항목들
        if doc_type == "제품설명회 시행 신청서":
            multi_item_fields = {
                "직원팀명": "직원팀명", 
                "팀명성명": "직원성명",
                "의료기관명": "의료기관명",
                "보건의료전문가성명": "보건의료전문가성명"
            }
        else:
            multi_item_fields = {
                "참석직원팀명": "직원팀명", 
                "참석직원성명": "직원성명",
                "참석의료기관명": "의료기관명",
                "참석보건의료전문가성명": "보건의료전문가성명"
            }
        
        # 문서에서 실제로 사용되는 플레이스홀더 번호 범위를 동적으로 찾기
        max_placeholders = self._find_max_placeholder_numbers(doc, multi_item_fields.keys())
        
        # 일반 플레이스홀더 치환 (문단)
        for paragraph in doc.paragraphs:
            self._replace_in_text_element(paragraph, filled_data, multi_item_fields, max_placeholders)
        
        # 테이블 내용 치환
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_text_element(paragraph, filled_data, multi_item_fields, max_placeholders)

    def _find_max_placeholder_numbers(self, doc, field_keys):
        """
        문서에서 실제로 사용되는 최대 플레이스홀더 번호를 찾습니다.
        
        Args:
            doc: DOCX 문서 객체
            field_keys (list): 확인할 필드 키 목록
        
        Returns:
            dict: 각 필드별 최대 번호 딕셔너리
        """
        max_numbers = {}
        
        # 모든 텍스트 수집
        all_text = ""
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_text += paragraph.text + "\n"
        
        # 각 필드별 최대 번호 찾기
        for field_key in field_keys:
            pattern = rf"{field_key}항목내용(\d+)"
            numbers = re.findall(pattern, all_text)
            if numbers:
                max_numbers[field_key] = max(int(n) for n in numbers)
            else:
                max_numbers[field_key] = 0
        
        return max_numbers

    def _replace_in_text_element(self, text_element, filled_data, multi_item_fields, max_placeholders):
        """
        텍스트 요소에서 플레이스홀더를 치환합니다.
        
        Args:
            text_element: DOCX 텍스트 요소 (paragraph)
            filled_data (dict): 치환할 데이터
            multi_item_fields (dict): 다중 항목 필드 매핑
            max_placeholders (dict): 최대 플레이스홀더 번호
        
        Returns:
            None: 텍스트 요소를 직접 수정
        """
        
        # 모든 치환 작업을 수집
        replacements = {}
        
        # 일반 필드 처리  
        for key, value in filled_data.items():
            if key not in multi_item_fields.values():
                # 지급내역은 특별 처리
                if key == "지급내역":
                    placeholder = "제품설명회지급내역항목내용"
                    replacement_value = str(value) if value else ""
                    replacements[placeholder] = replacement_value
                # 개별 예산 필드들 처리
                elif key in ["1인금액", "금액", "메뉴", "주류"]:
                    placeholder = f"{key}항목내용"
                    replacement_value = str(value) if value else ""
                    replacements[placeholder] = replacement_value
                else:
                    placeholder = f"{key}항목내용"
                    replacement_value = str(value) if value else ""
                    replacements[placeholder] = replacement_value
        
        # 다중 항목 필드 처리
        for field_key, data_key in multi_item_fields.items():
            value = filled_data.get(data_key, "")
            # 콤마로 분리하여 리스트로 변환
            items = [item.strip() for item in str(value).split(',')] if value else []
            
            # 동적으로 찾은 최대 번호까지 처리
            max_num = max_placeholders.get(field_key, 0)
            for i in range(1, max_num + 1):
                placeholder = f"{field_key}항목내용{i}"
                replacement_value = items[i-1] if i-1 < len(items) else ""
                replacements[placeholder] = replacement_value
        
        # 템플릿에 있는 추가 플레이스홀더들 처리
        additional_placeholders = [
            "PM참석항목내용", "구분항목내용", "일시항목내용", "장소항목내용", 
            "제품명항목내용", "제품설명회시행목적항목내용", "제품설명회주요내용항목내용", 
            "참석인원항목내용", "방문일항목내용"
        ]
        
        for placeholder in additional_placeholders:
            if placeholder not in replacements:
                # 해당하는 데이터 키 찾기
                data_key = placeholder.replace("항목내용", "")
                # 특별한 매핑 처리
                if placeholder == "방문일항목내용":
                    data_key = "방문날짜"
                replacement_value = str(filled_data.get(data_key, ""))
                replacements[placeholder] = replacement_value
        
        # run 단위로 포맷팅을 유지하면서 치환
        self._replace_text_preserving_format(text_element, replacements)

    def _replace_text_preserving_format(self, paragraph, replacements):
        """
        포맷팅을 유지하면서 텍스트를 치환합니다.
        
        Args:
            paragraph: DOCX 문단 객체
            replacements (dict): 치환할 텍스트 매핑
        
        Returns:
            None: 문단 객체를 직접 수정
        """
        if not replacements:
            return
        
        # 모든 run에서 텍스트를 수집
        full_text = ""
        run_texts = []
        
        for run in paragraph.runs:
            run_text = run.text
            run_texts.append(run_text)
            full_text += run_text
        
        # 치환 작업 수행
        modified_text = full_text
        for placeholder, replacement in replacements.items():
            if placeholder in modified_text:
                # 특별 처리: 금액항목내용이 1인금액항목내용의 일부인지 확인
                if placeholder == "금액항목내용":
                    if "1인금액항목내용" not in modified_text:
                        modified_text = modified_text.replace(placeholder, replacement)
                    else:
                        # 1인금액항목내용이 아닌 금액항목내용만 매칭하는 패턴
                        pattern = r'(?<!1인)금액항목내용'
                        modified_text = re.sub(pattern, replacement, modified_text)
                else:
                    modified_text = modified_text.replace(placeholder, replacement)
        
        # 텍스트가 변경되었을 때만 처리
        if modified_text != full_text:
            # 모든 기존 run 제거
            for run in paragraph.runs[:]:
                run._element.getparent().remove(run._element)
            
            # 새로운 run으로 변경된 텍스트 추가
            paragraph.add_run(modified_text)   
        


    def run(self, request_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        주요 실행 함수 - 문서 생성 요청을 처리합니다.
        
        Args:
            request_data (dict): 프론트엔드에서 받은 요청 데이터
                - type (str): 문서 타입
                - content (dict): 문서 내용 데이터
        
        Returns:
            dict: 처리 결과를 포함한 응답
                - status (str): 'success' or 'violation'
                - message (str): 결과 메시지
                - document_type (str): 문서 타입
                - violations (str, optional): 위반 내용 (status가 'violation'인 경우)
                - file_path (str, optional): 생성된 문서 경로 (status가 'success'인 경우)
                - document_content (str, optional): 생성된 문서 내용 (status가 'success'인 경우)
        """
        # 1. input_data 함수를 사용해서 type과 content를 분류
        document_type, content_data = self.input_data(request_data)

        print(f'document_type : {document_type}')
        print(f'content_data : {content_data}')
        print('='*30)
        
        # 2. parse_user_input으로 LLM을 통한 데이터 파싱
        parsed_data = self.parse_user_input(document_type, content_data)
        print(f'parsed_data : {parsed_data}')
        print('='*30)
        
        # 3. policy_violation_content_extraction 함수를 사용하여 검토할 내용을 추출
        content_to_check = self.policy_violation_content_extraction(document_type, parsed_data)
        print(f'content_to_check : {content_to_check}')
        print('='*30)

        # 3. check_policy_violation 함수를 사용하여 정책 위반 여부 확인
        result = self.check_policy_violation(content_to_check)
        
        # 4. 결과에 따라 처리
        if result == "OK" or "관련 규정 정보를 찾을 수 없습니다" in result:
            # 위반사항이 없거나 관련 규정을 찾을 수 없는 경우 문서 생성
            file_path = self.create_document(document_type, parsed_data)
            
            if file_path:
                # 생성된 문서의 텍스트 내용 추출
                document_content = self._extract_document_content(file_path)
                
                return {
                    "status": "success",
                    "message": "문서가 성공적으로 생성되었습니다.",
                    "document_type": document_type,
                    "file_path": file_path,
                    "parsed_data": parsed_data,  # 파싱 결과 추가
                    "document_content": document_content  # 문서 내용 추가
                }
            else:
                return {
                    "status": "error",
                    "message": "문서 생성 중 오류가 발생했습니다.",
                    "document_type": document_type
                }
        else:
            # 위반사항이 있는 경우 출력
            return {
                "status": "violation",
                "message": "정책 위반 사항이 발견되었습니다.",
                "violations": result,
                "document_type": document_type
            }
    
    def _extract_document_content(self, file_path: str) -> str:
        """
        생성된 DOCX 파일에서 텍스트 내용을 추출합니다.
        
        Args:
            file_path (str): DOCX 파일 경로
            
        Returns:
            str: 문서의 텍스트 내용
        """
        try:
            from docx import Document
            doc = Document(file_path)
            
            # 문서의 모든 텍스트 추출
            full_text = []
            
            # 단락 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # 테이블 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
            
        except Exception as e:
            print(f"ERROR: 문서 내용 추출 중 오류: {e}")
            return "문서 내용을 읽을 수 없습니다."


if __name__ == "__main__":    
    test_body1 = {
        "type" : "영업방문결과보고서",
        "content" : {
            "방문일" : "250808",
            "병원명" : "죽은사람은못살리는의원",
            '고객사개요' :'이번에 새로 오픈한 가정의학과로 사용 약품에 대해 많은 논의가 필요해보이는 잠재력이 있는 고객', 
            '프로젝트개요' :'신규고객 유치로 자사에서 납품하는 의약품 소개', 
            '방문및협의내용' : '자사 약품 인 로바스로 제품 소개', 
            '향후계획및일정' :'7월 27일에 다시 방문하여 자사 판촉물 전달과 약품별 로얄티 금액 소개 및 협상 예정', 
            '협조사항및공유사항':'재방문 전에 유미가정의학과에 전달할 자사 판촉물 1개 지급 요망'
        }
    }

    test_body2 = {
        "type" : "영업방문결과보고서",
        "content" : {
            "방문일" : "250808",
            "병원명" : "죽은사람은못살리는의원",
            '고객사개요' :'문서생성확인용', 
            '프로젝트개요' :'문서생성확인용', 
            '방문및협의내용' : '문서생성확인용', 
            '향후계획및일정' :'문서생성확인용', 
            '협조사항및공유사항':'문서생성확인용'
        }
    }

    test_body3 = {
        "type" : "제품설명회시행결과보고서",
        "content" : {
            "구분": "제품설명회",
            "PM참석": "김도윤",
            "일시": "25년 7월 25일",
            "장소": "코엑스 B홀",
            "제품명": "텐텐",
            "참석인원": "5명",
            "제품설명회시행목적": "제품 리뉴얼 소개",
            "제품설명회주요내용": "기존 제품의 문제점과 리뉴얼 되면서 바뀐점과 영양성분 소개",
            "지급내역": "",
            "금액": "8만원",
            "메뉴": "보쌈",
            "주류": "소주1병, 맥주2병",
            "1인금액": "1만 6천원",
            "참석직원": [['영업팀', '손현성'], ['영업팀','이용규'], ['영업팀','손영식']],
            "참석의료전문가": [['서울대학병원', '허한결'], ['연세대학병원', '최문영']]
        }
    }

    test_body4 = {
        'type' : '영업방문결과보고서',
        'content' : {
            '방문제목': '유미가정의학과 신약 홍보', 
            '방문일': '2025-07-25', 
            '병원명': '유미가정의학과', 
            '지역구': '강남구', 
            '원장명': '손영식', 
            '원장연락처': '010-1234-5678', 
            '담당자성명': '손현성', 
            '담당자부서': '영업팀', 
            '담당자연락처': '010-8765-4321', 
            '지점': '강남팀', 
            '지점연락처': '02-123-4567', 
            '고객사개요': '이번에 새로 오픈한 가정의학과로 사용 약품에 대해 많은 논의가 필요해보이는 잠재력이 있는 고객', 
            '프로젝트개요': '신규고객 유치로 자사에서 납품하는 의약품 소개', 
            '방문및협의내용': '자사 약품인 로바스를 제품 소개하였습니다.', 
            '향후계획및일정': '2023년 7월 27일에 다시 방문하여 자사 판촉물을 전달하고 약품별 로얄티 금액 소개 및 협상을 진행할 예정입니다.', 
            '협조사항및공유사항': '재방문 전에 유미가정의학과에 전달할 자사 판촉물 1개 지급을 요청드립니다.'
        }
    }

    # 인스턴스 생성 및 테스트 실행
    doc_func = CreateDocumentFunction()
    
    # 제품설명회시행결과보고서 테스트
    print("=== 제품설명회시행결과보고서 테스트 ===")
    result = doc_func.run(test_body4)
    print(result)
    