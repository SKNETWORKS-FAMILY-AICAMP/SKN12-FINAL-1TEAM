from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage
from langchain_core.prompts import ChatPromptTemplate
from typing import TypedDict, List, Optional
import json
import yaml
import time
import re
import os
import uuid
from pathlib import Path
from dotenv import load_dotenv
from docx import Document

# 외부 도구 임포트
import sys
# 상위 디렉토리를 sys.path에 추가하여 tools 모듈에 접근
sys.path.append(str(Path(__file__).parent.parent))
from tools.common_tools import check_policy_violation, separate_document_type_and_content

load_dotenv()

class State(TypedDict):
    messages: List[HumanMessage]
    doc_type: Optional[str]
    template_content: Optional[str]
    filled_data: Optional[dict]
    violation: Optional[str]
    final_doc: Optional[str]
    classification_failed: Optional[bool]  # 분류 실패 플래그
    skip_verification: Optional[bool]  # 검증 건너뛰기 플래그
    end_process: Optional[bool]
    parse_retry_count: Optional[int]
    parse_failed: Optional[bool]
    user_reply: Optional[str]  # 휴먼인더루프용 사용자 입력
    verification_reply: Optional[str]  # 분류 검증용 사용자 입력
    verification_result: Optional[str]  # 긍정/부정 분류 결과
    user_content: Optional[str]  # 문서 내용 (separate_document_type_and_content에서 추출)
    skip_ask_fields: Optional[bool]  # ask_required_fields 스킵 플래그

class CreateDocumentAgent:
    """통합 문서 작성 에이전트 - 분류부터 생성까지"""
    
    def __init__(self, model_name: str = "gpt-4o-mini", temperature: float = 0.7, api_mode: bool = False):
        """
        CreateDocumentAgent 초기화
        
        Args:
            model_name: 기본 LLM 모델명
            temperature: LLM 온도 설정
            api_mode: API 모드 활성화 여부. True일 경우 input() 대신 인터럽트 정보 반환
        """
        self.model_name = model_name
        self.temperature = temperature
        self.api_mode = api_mode or os.getenv("NO_INPUT_MODE", "false").lower() == "true"
        
        # LLM 초기화
        self.llm = ChatOpenAI(
            model=self.model_name, 
            temperature=self.temperature
        )
        
        # YAML 파일에서 템플릿 로드
        self.doc_prompts = self._load_templates()
        
        # 그래프 초기화
        self.app = self._build_graph()
    
    def _load_templates(self):
        """
        YAML 파일에서 문서 템플릿 및 프롬프트 정보를 로드합니다.
        
        템플릿 파일 구조:
        - templates.yaml 파일에서 각 문서 타입별 템플릿 정보 로드
        - 각 문서 타입마다 input_prompt, choan_system_prompt, choan_fallback_fields 포함
        
        Returns:
            dict: 문서 타입별 템플릿 정보 딕셔너리
                 예: {"영업방문 결과보고서": {"input_prompt": "...", "choan_system_prompt": "...", ...}}
                 로드 실패 시 빈 딕셔너리 반환
        
        Raises:
            Exception: 파일 읽기 실패 시 예외 처리하고 빈 딕셔너리 반환
        """
        try:
            # 현재 스크립트와 같은 디렉토리에서 templates.yaml 파일 찾기
            current_dir = Path(__file__).parent
            template_path = current_dir / "templates.yaml"
            
            # 템플릿 파일 존재 여부 확인
            if not template_path.exists():
                print(f"[WARNING] 템플릿 파일을 찾을 수 없습니다: {template_path}")
                return {}
            
            # YAML 파일 읽기 및 파싱
            with open(template_path, 'r', encoding='utf-8') as file:
                data = yaml.safe_load(file)
                return data.get('templates', {})
                
        except Exception as e:
            print(f"[ERROR] 템플릿 로드 중 오류 발생: {e}")
            return {}
    

    def classify_doc_type(self, state: State) -> State:
        """
        separate_document_type_and_content 툴을 사용하여 사용자 요청에서 문서 타입과 내용을 분리하고,
        분리된 문서 타입을 LLM으로 분류합니다.
        
        Args:
            state (State): 현재 워크플로우 상태
                - messages: 사용자 메시지 리스트 (최신 메시지에서 요청 내용 추출)
        
        Returns:
            State: 업데이트된 상태
                - doc_type: 분류된 문서 타입 또는 "분류 실패"
                - user_content: 문서에 들어갈 내용
                - skip_ask_fields: 내용이 있으면 True, 없으면 False
        """
        # 사용자 메시지에서 최신 요청 내용 추출
        user_message = state["messages"][-1].content
        
        try:
            # 1단계: separate_document_type_and_content 툴로 문서 타입과 내용 분리
            print("[SEARCH] 문서 타입과 내용 분리 중...")
            separation_result = separate_document_type_and_content.invoke({"user_input": user_message})
            
            # JSON 파싱
            import json
            separated_data = json.loads(separation_result)
            document_type_text = separated_data.get("document_type", "")
            content_text = separated_data.get("content", "")
            
            print(f"[INFO] 분리된 문서 타입: '{document_type_text}'")
            print(f"[INFO] 분리된 내용: '{content_text[:50]}...' (길이: {len(content_text)})")
            
            # 상태에 내용 저장
            state["user_content"] = content_text
            state["skip_ask_fields"] = bool(content_text.strip())  # 내용이 있으면 True
            
            # 2단계: 분리된 문서 타입을 LLM으로 분류
            if document_type_text.strip():
                classification_input = document_type_text
            else:
                classification_input = user_message  # 문서 타입이 분리되지 않았으면 전체 메시지 사용
            
            classification_prompt = ChatPromptTemplate.from_messages([
                ("system", """
사용자의 요청을 분석하여 다음 문서 타입 중 하나로 분류해주세요:
1. 영업방문 결과보고서 - 고객 방문, 영업 활동 관련
2. 제품설명회 시행 신청서 - 제품설명회 진행 계획, 신청 관련
3. 제품설명회 시행 결과보고서 - 제품설명회 완료 후 결과 보고 관련

반드시 위 3가지 중 하나의 정확한 문서 타입 이름만 응답해주세요.
앞에 숫자는 제거하고 문서명만 출력하세요.
                """),
                ("human", "{user_request}")
            ])
            
            # LLM을 통한 문서 타입 분류 실행
            response = self.llm.invoke(classification_prompt.format_messages(user_request=classification_input))
            content = response.content
            
            # 응답 내용을 문자열로 정규화
            if isinstance(content, str):
                doc_type = content.strip()
            else:
                doc_type = str(content).strip()
                
            # 분류 결과를 상태에 저장
            state["doc_type"] = doc_type
            print(f"[INFO] LLM 문서 타입 분류: {doc_type}")
            print(f"[INFO] ask_required_fields 스킵 여부: {state['skip_ask_fields']}")
            
        except Exception as e:
            # 처리 실패 시 예외 처리
            print(f"[WARNING] 문서 분류 및 분리 실패: {e}")
            state["doc_type"] = "분류 실패"
            state["user_content"] = ""
            state["skip_ask_fields"] = False
        
        return state

    def validate_doc_type(self, state: State) -> State:
        """
        분류된 문서 타입의 유효성을 검증하고 해당 템플릿 정보를 로드합니다.
        
        이 함수는 classify_doc_type에서 분류된 문서 타입이 시스템에서 지원하는
        유효한 타입인지 확인하고, 유효한 경우 해당 템플릿 정보를 state에 추가합니다.
        
        Args:
            state (State): 현재 워크플로우 상태
                - doc_type: 분류된 문서 타입 (classify_doc_type에서 설정)
        
        Returns:
            State: 업데이트된 상태
                성공 시:
                - classification_failed: False
                - template_content: 해당 문서 타입의 input_prompt 템플릿
                실패 시:
                - classification_failed: True  
                - skip_verification: True (검증 단계 건너뛰기)
        """
        # 상태에서 분류된 문서 타입 추출
        doc_type = state.get("doc_type", "")
        # 시스템에서 지원하는 유효한 문서 타입 목록
        valid_types = ["영업방문 결과보고서", "제품설명회 시행 신청서", "제품설명회 시행 결과보고서"]
        
        print(f"[SEARCH] 문서 타입 검증 중: '{doc_type}'")
        print(f"[SEARCH] 유효한 타입 목록: {valid_types}")
        
        # 유효한 문서 타입인지 확인
        if doc_type in valid_types:
            print(f"[SUCCESS] 유효한 문서 타입: {doc_type}")
            
            # 분류된 문서 타입에 맞는 템플릿을 state에 추가
            if doc_type in self.doc_prompts:
                state["template_content"] = self.doc_prompts[doc_type]["input_prompt"]
                print(f"[INFO] 템플릿 추가 완료: {doc_type}")
            
            # 분류 성공 플래그 설정
            state["classification_failed"] = False
            print(f"[SEARCH] classification_failed 설정: False")
            return state
        else:
            print(f"[ERROR] 유효하지 않은 문서 타입: '{doc_type}'")
            print("[INFO] 자동 분류 실패 - 수동 선택으로 직접 이동합니다.")
            
            # 분류 실패 플래그 설정 - 명시적으로 True 설정
            state["classification_failed"] = True
            print(f"[SEARCH] classification_failed 설정: True")
            
            # 추가 보안: 검증 건너뛰기 플래그도 설정 (verify_classification 노드 건너뛰기)
            state["skip_verification"] = True
            print(f"[SEARCH] skip_verification 설정: True")
            
            return state

    def process_verification_response(self, state: State) -> State:
        """
        사용자의 분류 검증 응답을 LLM을 통해 분석하여 긍정/부정을 판단합니다.
        
        이 함수는 사용자가 입력한 자연어 응답을 LLM을 통해 분석하여
        분류 결과에 동의하는지(긍정) 또는 반대하는지(부정)를 판단합니다.
        
        Args:
            state (State): 현재 워크플로우 상태
                - verification_reply: 외부에서 설정된 사용자 검증 응답
                - messages: 사용자 메시지 리스트 (검증 응답을 메시지에 추가)
        
        Returns:
            State: 업데이트된 상태
                - verification_result: "긍정", "부정", "불명확", "오류" 중 하나
        """
        # verification_reply에서 사용자 응답 추출
        user_response = state.get("verification_reply", "")
        
        if not user_response:
            # 메시지에서 최신 응답 추출 (폴백)
            if state.get("messages"):
                user_response = state["messages"][-1].content
            else:
                print("[WARNING] 사용자 검증 응답이 없습니다.")
                return state
        
        # 사용자 검증 응답을 메시지에 추가
        if state.get("messages") is None:
            state["messages"] = []
        state["messages"].append(HumanMessage(content=user_response))
        
        # verification_reply 정리
        state["verification_reply"] = None
        
        # LLM을 통한 응답 분석을 위한 프롬프트 구성
        verification_prompt = ChatPromptTemplate.from_messages([
            ("system", """
사용자의 응답을 분석하여 긍정인지 부정인지 판단해주세요.

긍정적 응답: YES, 네, 맞습니다, 맞아요, 정확합니다, 올바릅니다, 그렇습니다, 동의합니다 등
부정적 응답: NO, 아니요, 틀렸습니다, 틀려요, 잘못됐습니다, 다릅니다, 아닙니다 등

응답 형식: "긍정" 또는 "부정"만 출력해주세요.
            """),
            ("human", "{user_response}")
        ])
        
        try:
            # LLM을 통한 응답 분석 실행
            response = self.llm.invoke(verification_prompt.format_messages(user_response=user_response))
            content = response.content.strip()
            
            # 분석 결과에 따른 상태 업데이트
            if "긍정" in content:
                state["verification_result"] = "긍정"
                print(f"[SUCCESS] 분류 검증 결과: 긍정 - 기존 분류를 유지합니다.")
            elif "부정" in content:
                state["verification_result"] = "부정"
                print(f"[ERROR] 분류 검증 결과: 부정 - 새로운 문서 타입을 선택해주세요.")
            else:
                # 분석 결과가 명확하지 않은 경우
                print(f"[WARNING] 검증 응답 분석 실패: {content}")
                state["verification_result"] = "불명확"
                
        except Exception as e:
            # LLM 호출 실패 시 예외 처리
            print(f"[WARNING] 검증 응답 분석 중 오류: {e}")
            state["verification_result"] = "오류"
        
        return state

    def process_manual_doc_type_selection(self, state: State) -> State:
        """
        사용자가 선택한 문서 타입을 처리합니다.
        
        Args:
            state (State): verification_reply 필드에 사용자 선택 포함
        
        Returns:
            State: doc_type, template_content, end_process 업데이트된 상태
        """
        # verification_reply에서 사용자 선택 추출
        user_selection = state.get("verification_reply", "")
        
        if not user_selection:
            # 메시지에서 최신 선택 추출 (폴백)
            if state.get("messages"):
                user_selection = state["messages"][-1].content
                # 만약 이미 메시지에 추가된 것이면 중복 추가하지 않음
            else:
                print("[WARNING] 사용자 문서 타입 선택이 없습니다.")
                return state
        else:
            # verification_reply가 있으면 메시지에 추가
            if state.get("messages") is None:
                state["messages"] = []
            state["messages"].append(HumanMessage(content=user_selection))
        
        # verification_reply 정리
        state["verification_reply"] = None
        
        user_selection = user_selection.strip()
        
        # 문서 타입 매핑
        doc_type_mapping = {
            "1": "영업방문 결과보고서",
            "2": "제품설명회 시행 신청서", 
            "3": "제품설명회 시행 결과보고서",
            "4": "종료",
            "영업방문 결과보고서": "영업방문 결과보고서",
            "제품설명회 시행 신청서": "제품설명회 시행 신청서",
            "제품설명회 시행 결과보고서": "제품설명회 시행 결과보고서",
            "종료": "종료"
        }
        
        selected_doc_type = doc_type_mapping.get(user_selection)
        
        if selected_doc_type == "종료":
            print("[INFO] 사용자가 종료를 선택했습니다.")
            state["end_process"] = True
            return state
        elif selected_doc_type:
            print(f"[SUCCESS] 선택된 문서 타입: {selected_doc_type}")
            state["doc_type"] = selected_doc_type
            
            # 선택된 문서 타입에 맞는 템플릿 설정
            if selected_doc_type in self.doc_prompts:
                state["template_content"] = self.doc_prompts[selected_doc_type]["input_prompt"]
                print(f"[INFO] 템플릿 업데이트 완료: {selected_doc_type}")
            
            return state
        else:
            print(f"[ERROR] 유효하지 않은 선택: '{user_selection}'")
            print("올바른 번호(1-4) 또는 정확한 문서명을 입력해주세요.")
            print("예시: '1' 또는 '영업방문 결과보고서'")
            print("[INFO] 문서 타입 선택 실패로 인해 에이전트를 종료합니다.")
            # 선택 실패 시 에이전트 종료
            state["end_process"] = True
            return state

    def receive_user_input(self, state: State) -> State:
        """
        외부에서 입력을 받아 재개되는 노드
        
        Args:
            state (State): user_reply 필드 포함
        
        Returns:
            State: messages에 새 메시지 추가, user_reply를 None으로 초기화
        """
        user_reply = state.get("user_reply", "")
        
        if user_reply:
            print(f"[SUCCESS] 사용자 입력 수신됨: {user_reply[:50]}...")
            # 사용자 입력을 메시지에 추가
            if state.get("messages") is None:
                state["messages"] = []
            state["messages"].append(HumanMessage(content=user_reply))
            # user_reply 플래그 제거
            state["user_reply"] = None
        else:
            print("[WARNING] 사용자 입력이 없습니다.")
            
        return state

    def parse_user_input(self, state: State) -> State:
        """
        사용자 입력 또는 user_content를 LLM을 통해 구조화된 JSON 데이터로 파싱합니다.
        
        Args:
            state (State): messages, doc_type, parse_retry_count, user_content 필드 포함
        
        Returns:
            State: filled_data, parse_failed 업데이트된 상태
        """
        # 내용이 이미 있으면 user_content 사용, 없으면 최신 메시지 사용
        user_content = state.get("user_content", "")
        if user_content.strip():
            user_input = user_content
            print(f"[INFO] 미리 분리된 내용 사용: '{user_input[:50]}...'")
        else:
            user_input = str(state["messages"][-1].content)
            print(f"[INFO] 사용자 입력 사용: '{user_input[:50]}...'")
            
        doc_type = state["doc_type"]
        response = None

        if state.get("parse_retry_count") is None:
            state["parse_retry_count"] = 0

        system_prompt = self.doc_prompts[doc_type]["choan_system_prompt"]
        if not system_prompt:
            raise ValueError(f"문서 타입에 대한 시스템 프롬프트가 없습니다: {doc_type}")

        # 중괄호 이스케이프 처리
        escaped_input = user_input.replace("{", "{{").replace("}", "}}")

        parsing_prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{user_input}")
        ])

        try:
            formatted_messages = parsing_prompt.format_messages(user_input=escaped_input)
            # API 모드에서 규정 위반이 없을 때만 상세 로그 출력
            violation = state.get("violation", "")
            if not self.api_mode or violation == "OK":
                print("[INFO] LLM에 전달된 메시지:")
                for m in formatted_messages:
                    print(f"[{m.type.upper()}] {m.content[:200]}...")

            response = self.llm.invoke(formatted_messages)

            content = response.content
            json_str = content if isinstance(content, str) else str(content)
            
            # API 모드에서 규정 위반 시에는 상세 로그 출력 생략
            violation = state.get("violation", "")
            if not self.api_mode or violation == "OK":
                print(f"\n[SEARCH] LLM 응답 내용:\n{json_str}")

            if "{" in json_str and "}" in json_str:
                start = json_str.find("{")
                end = json_str.rfind("}") + 1
                clean_json = json_str[start:end]
                
                if not self.api_mode or violation == "OK":
                    print(f"\n[SEARCH] 추출된 JSON:\n{clean_json}")

                try:
                    parsed_data = json.loads(clean_json)
                    state["filled_data"] = parsed_data
                    state["parse_failed"] = False
                    
                    # API 모드에서 규정 위반 시에는 상세 파싱 결과 출력 생략
                    violation = state.get("violation", "")
                    if not self.api_mode or violation == "OK":
                        print("[SUCCESS] 파싱 성공:", parsed_data)
                    else:
                        print("[SUCCESS] 파싱 성공 (상세 결과 생략)")
                except json.JSONDecodeError as json_error:
                    print(f"[ERROR] JSON 파싱 오류: {json_error}")
                    print(f"파싱 시도한 JSON: {repr(clean_json)}")
                    raise json_error
            else:
                raise ValueError("구조화된 JSON 형식을 찾을 수 없음")

        except Exception as e:
            print("\n[WARNING] 예외 발생!")
            if response:
                print("응답 내용:")
                print(response)
            else:
                print("[WARNING] response 객체가 존재하지 않습니다.")
            print(f"[WARNING] 예외 메시지: {e}")

            retry_count = state.get("parse_retry_count", 0) + 1
            state["parse_retry_count"] = retry_count

            if retry_count >= 3:
                print("WARNING: 파싱 재시도 초과. 기본값 사용.")
                fallback_data = self.doc_prompts[doc_type]["choan_fallback_fields"]
                state["filled_data"] = fallback_data
            else:
                print(f"RETRY: 재시도 {retry_count}/3")
                state["parse_failed"] = True

        return state

    def check_user_input_policy(self, state: State) -> State:
        """
        사용자 입력 텍스트를 직접 규정 위반 검사합니다.
        
        Args:
            state (State): user_reply, messages, user_content 필드 포함
        
        Returns:
            State: violation 필드 업데이트된 상태
        """
        # user_reply에서 사용자 입력 추출
        user_reply = state.get("user_reply", "")
        
        if user_reply:
            # 사용자 입력을 메시지에 추가
            if state.get("messages") is None:
                state["messages"] = []
            state["messages"].append(HumanMessage(content=user_reply))
            
            # user_reply 정리
            state["user_reply"] = None
            
            input_text = user_reply
            print(f"POLICY_CHECK: 사용자 입력으로 규정 검사: '{input_text[:50]}...'")
        else:
            # 내용이 이미 있으면 user_content 사용, 없으면 최신 메시지 사용
            user_content = state.get("user_content", "")
            if user_content.strip():
                input_text = user_content
                print(f"POLICY_CHECK: 미리 분리된 내용으로 규정 검사: '{input_text[:50]}...'")
            else:
                if state.get("messages"):
                    input_text = str(state["messages"][-1].content)
                    print(f"POLICY_CHECK: 사용자 입력으로 규정 검사: '{input_text[:50]}...'")
                else:
                    print("WARNING: 검사할 입력이 없습니다.")
                    return state
        
        try:
            print("POLICY_CHECK: 규정 위반 검사 시작...")
            violation_result = check_policy_violation.invoke({"content": input_text})
            
            state["violation"] = violation_result
            
            if violation_result == "OK":
                print("SUCCESS: 규정 위반 없음 - 파싱 단계로 진행")
            else:
                print(f"WARNING: 규정 위반 사항 발견: {violation_result}")
            
        except Exception as e:
            print(f"ERROR: 규정 검사 중 오류 발생: {e}")
            import traceback
            traceback.print_exc()
            state["violation"] = f"규정 검사 오류: {str(e)}"
        
        return state

    def inform_violation(self, state: State) -> State:
        """
        규정 위반이 발견되었을 때 위반 내용을 안내하고 종료합니다.
        
        Args:
            state (State): violation 필드 포함
        
        Returns:
            State: final_doc=None, end_process=True로 설정된 상태
        """
        violation = state["violation"]
        
        if self.api_mode:
            # API 모드에서는 간단하게 위반 내용만 출력
            print(f"\n[위반 내용]\n{violation}")
        else:
            # 콘솔 모드에서는 상세한 안내
            print(f"\nALERT: 규정 위반 사항 발견!")
            print("=" * 60)
        
            # check_policy_violation 결과를 직접 출력
            print("VIOLATION_RESULT: 규정 위반 검토 결과:")
            print(violation)
            
            print("=" * 60)
            print("ERROR: 규정 위반 사항으로 인해 문서 생성을 중단합니다.")
            print("INFO: 위반 내용을 확인하고 내용을 수정한 후 다시 시도해주세요.")
        
        # 위반 사항을 state에 저장
        state["final_doc"] = None  # 문서 생성 실패 표시
        state["end_process"] = True  # 프로세스 종료 표시
        state["error_type"] = "policy_violation"  # 오류 타입 명시
        
        return state
    
    def create_choan_document(self, state: State) -> State:
        """
        파싱된 데이터를 기반으로 초안 문서를 생성하고 docx 파일로 저장합니다.
        규정 위반이 있는 경우 파일 생성을 차단합니다.
        
        Args:
            state (State): doc_type, filled_data, violation 필드 포함
        
        Returns:
            State: final_doc 필드 업데이트된 상태 (파일 경로 또는 None)
        """
        doc_type = state["doc_type"]
        filled_data = state["filled_data"]
        violation = state.get("violation", "")
        
        # API 모드에서는 여기까지 도달하지 않음 (규정 위반 시 inform_violation에서 처리됨)
        # 콘솔 모드에서만 실행
        if violation and violation.strip() != "OK":
            print("\n[ALERT] 규정 위반으로 인해 파일 생성이 차단되었습니다!")
            print("[INFO] 규정 위반으로 문서 파일은 생성되지 않습니다.")
            
            # 위반 내용 표시
            print(f"\n[위반 내용] {violation}")
            
            state["final_doc"] = None
            state["violation_blocked"] = True
            return state
        
        # 문서 타입에 따른 템플릿 파일 매핑
        template_mapping = {
            "영업방문 결과보고서": "영업방문 결과보고서(템플릿형).docx",
            "제품설명회 시행 신청서": "제품설명회 시행 신청서(템플릿형).docx",
            "제품설명회 시행 결과보고서": "제품설명회 시행 결과보고서(템플릿형).docx"
        }
        template_filename = template_mapping.get(doc_type)
        if not template_filename:
            print(f"ERROR: 지원하지 않는 문서 타입: {doc_type}")
            state["final_doc"] = None
            return state
        
        # S3 폴더에서 템플릿 파일 경로 구성
        current_dir = Path(__file__).parent
        template_path = current_dir / "S3" / template_filename
        
        if not template_path.exists():
            print(f"ERROR: 템플릿 파일을 찾을 수 없습니다: {template_path}")
            state["final_doc"] = None
            return state
        
        try:
            # 템플릿 파일 읽기
            print(f"LOADING: 템플릿 파일 로딩: {template_filename}")
            doc = Document(str(template_path))
            
            print(f"PROCESSING: 템플릿 플레이스홀더 치환 중...")
            
            # 양식을 유지하면서 플레이스홀더만 치환
            self._replace_placeholders_in_document(doc, filled_data, doc_type)
            
            # agent_result_folder 디렉토리 생성
            result_folder = current_dir / "agent_result_folder"
            result_folder.mkdir(exist_ok=True)
            
            # 완성된 문서 저장
            today_date = time.strftime('%Y%m%d')
            doc_type_no_space = doc_type.replace(" ", "")
            output_filename = f"{doc_type_no_space}_{today_date}.docx"
            output_path = result_folder / output_filename
            doc.save(str(output_path))
            
            state["final_doc"] = str(output_path)
            
            print("SUCCESS: 문서 생성 및 저장 완료!")
            print(f"SAVED_PATH: 저장 경로: {output_path}")
            print("INFO: 템플릿 양식이 그대로 유지되면서 플레이스홀더만 치환되었습니다.")
            
        except Exception as e:
            print(f"ERROR: 문서 생성 중 오류 발생: {e}")
            import traceback
            traceback.print_exc()
            state["final_doc"] = None
        
        return state

    def _replace_placeholders_in_document(self, doc, filled_data, doc_type):
        """
        문서의 플레이스홀더를 실제 데이터로 치환합니다.
        
        Args:
            doc: DOCX 문서 객체
            filled_data (dict): 치환할 데이터
            doc_type (str): 문서 타입
        
        Returns:
            None: 문서 객체를 직접 수정
        """
        
        # 다중 항목 처리를 위한 특별 처리 필요 항목들
        if doc_type == "제품설명회 시행 신청서":
            multi_item_fields = {
                "직원팀명": "직원팀명", 
                "팀명성명": "직원성명",
                "의료기관명": "의료기관명",
                "보건의료전문가성명": "보건의료전문가성명"
            }
        else:
            multi_item_fields = {
                "참석직원팀명": "직원팀명", 
                "참석직원성명": "직원성명",
                "참석의료기관명": "의료기관명",
                "참석보건의료전문가성명": "보건의료전문가성명"
            }
        
        # 문서에서 실제로 사용되는 플레이스홀더 번호 범위를 동적으로 찾기
        max_placeholders = self._find_max_placeholder_numbers(doc, multi_item_fields.keys())
        
        # 일반 플레이스홀더 치환 (문단)
        for paragraph in doc.paragraphs:
            self._replace_in_text_element(paragraph, filled_data, multi_item_fields, max_placeholders)
        
        # 테이블 내용 치환
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_text_element(paragraph, filled_data, multi_item_fields, max_placeholders)

    def _find_max_placeholder_numbers(self, doc, field_keys):
        """
        문서에서 실제로 사용되는 최대 플레이스홀더 번호를 찾습니다.
        
        Args:
            doc: DOCX 문서 객체
            field_keys (list): 확인할 필드 키 목록
        
        Returns:
            dict: 각 필드별 최대 번호 딕셔너리
        """
        max_numbers = {}
        
        # 모든 텍스트 수집
        all_text = ""
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_text += paragraph.text + "\n"
        
        # 각 필드별 최대 번호 찾기
        for field_key in field_keys:
            pattern = rf"{field_key}항목내용(\d+)"
            numbers = re.findall(pattern, all_text)
            if numbers:
                max_numbers[field_key] = max(int(n) for n in numbers)
            else:
                max_numbers[field_key] = 0
        
        return max_numbers

    def _replace_in_text_element(self, text_element, filled_data, multi_item_fields, max_placeholders):
        """
        텍스트 요소에서 플레이스홀더를 치환합니다.
        
        Args:
            text_element: DOCX 텍스트 요소 (paragraph)
            filled_data (dict): 치환할 데이터
            multi_item_fields (dict): 다중 항목 필드 매핑
            max_placeholders (dict): 최대 플레이스홀더 번호
        
        Returns:
            None: 텍스트 요소를 직접 수정
        """
        
        # 모든 치환 작업을 수집
        replacements = {}
        
        # 일반 필드 처리  
        for key, value in filled_data.items():
            if key not in multi_item_fields.values():
                # 지급내역은 특별 처리
                if key == "지급내역":
                    placeholder = "제품설명회지급내역항목내용"
                    replacement_value = str(value) if value else ""
                    replacements[placeholder] = replacement_value
                # 개별 예산 필드들 처리
                elif key in ["1인금액", "금액", "메뉴", "주류"]:
                    placeholder = f"{key}항목내용"
                    replacement_value = str(value) if value else ""
                    replacements[placeholder] = replacement_value
                else:
                    placeholder = f"{key}항목내용"
                    replacement_value = str(value) if value else ""
                    replacements[placeholder] = replacement_value
        
        # 다중 항목 필드 처리
        for field_key, data_key in multi_item_fields.items():
            value = filled_data.get(data_key, "")
            # 콤마로 분리하여 리스트로 변환
            items = [item.strip() for item in str(value).split(',')] if value else []
            
            # 동적으로 찾은 최대 번호까지 처리
            max_num = max_placeholders.get(field_key, 0)
            for i in range(1, max_num + 1):
                placeholder = f"{field_key}항목내용{i}"
                replacement_value = items[i-1] if i-1 < len(items) else ""
                replacements[placeholder] = replacement_value
        
        # 템플릿에 있는 추가 플레이스홀더들 처리
        additional_placeholders = [
            "PM참석항목내용", "구분항목내용", "일시항목내용", "장소항목내용", 
            "제품명항목내용", "제품설명회시행목적항목내용", "제품설명회주요내용항목내용", 
            "참석인원항목내용", "방문일항목내용"
        ]
        
        for placeholder in additional_placeholders:
            if placeholder not in replacements:
                # 해당하는 데이터 키 찾기
                data_key = placeholder.replace("항목내용", "")
                # 특별한 매핑 처리
                if placeholder == "방문일항목내용":
                    data_key = "방문날짜"
                replacement_value = str(filled_data.get(data_key, ""))
                replacements[placeholder] = replacement_value
        
        # run 단위로 포맷팅을 유지하면서 치환
        self._replace_text_preserving_format(text_element, replacements)

    def _replace_text_preserving_format(self, paragraph, replacements):
        """
        포맷팅을 유지하면서 텍스트를 치환합니다.
        
        Args:
            paragraph: DOCX 문단 객체
            replacements (dict): 치환할 텍스트 매핑
        
        Returns:
            None: 문단 객체를 직접 수정
        """
        if not replacements:
            return
        
        # 모든 run에서 텍스트를 수집
        full_text = ""
        run_texts = []
        
        for run in paragraph.runs:
            run_text = run.text
            run_texts.append(run_text)
            full_text += run_text
        
        # 치환 작업 수행
        modified_text = full_text
        for placeholder, replacement in replacements.items():
            if placeholder in modified_text:
                # 특별 처리: 금액항목내용이 1인금액항목내용의 일부인지 확인
                if placeholder == "금액항목내용":
                    if "1인금액항목내용" not in modified_text:
                        modified_text = modified_text.replace(placeholder, replacement)
                    else:
                        # 1인금액항목내용이 아닌 금액항목내용만 매칭하는 패턴
                        pattern = r'(?<!1인)금액항목내용'
                        modified_text = re.sub(pattern, replacement, modified_text)
                else:
                    modified_text = modified_text.replace(placeholder, replacement)
        
        # 텍스트가 변경되었을 때만 처리
        if modified_text != full_text:
            # 모든 기존 run 제거
            for run in paragraph.runs[:]:
                run._element.getparent().remove(run._element)
            
            # 새로운 run으로 변경된 텍스트 추가
            paragraph.add_run(modified_text)

    def doc_type_validation_router(self, state: State) -> str:
        """
        문서 타입 유효성 검사 결과에 따라 다음 노드를 결정합니다.
        
        Args:
            state (State): classification_failed, skip_verification 필드 포함
        
        Returns:
            str: "process_manual_doc_type_selection" 또는 "process_verification_response"
        """
        classification_failed = state.get("classification_failed", False)
        skip_verification = state.get("skip_verification", False)
        doc_type = state.get("doc_type", "")
        
        print(f"ROUTER_STATUS: 라우터 상태 확인:")
        print(f"  - doc_type: '{doc_type}'")
        print(f"  - classification_failed: {classification_failed}")
        print(f"  - skip_verification: {skip_verification}")
        
        # 분류 실패이거나 검증 건너뛰기 플래그가 True인 경우
        if classification_failed or skip_verification:
            print(f"ROUTING: 라우팅 결정: process_manual_doc_type_selection (분류 실패)")
            return "process_manual_doc_type_selection"
        else:
            print(f"ROUTING: 라우팅 결정: process_verification_response (분류 성공)")
            return "process_verification_response"

    def verification_response_router(self, state: State) -> str:
        """
        검증 응답 결과에 따라 다음 노드를 결정합니다.
        내용이 있으면 ask_required_fields를 스킵하고 바로 check_user_input_policy으로 이동합니다.
        
        Args:
            state (State): verification_result, skip_ask_fields 필드 포함
        
        Returns:
            str: "receive_user_input", "check_user_input_policy", "process_manual_doc_type_selection", "process_verification_response" 중 하나
        """
        verification_result = state.get("verification_result", "")
        skip_ask_fields = state.get("skip_ask_fields", False)
        
        if verification_result == "긍정":
            # skip_ask_fields가 True면 이미 내용이 있으므로 바로 규정 검사로
            if skip_ask_fields:
                print("FLOW: 분륅 결과 승인 + 내용 있음 - 바로 check_user_input_policy로 이동")
                return "check_user_input_policy"
            else:
                print("FLOW: 분류 결과 승인 - receive_user_input로 이동")
                return "receive_user_input"
        elif verification_result == "부정":
            return "process_manual_doc_type_selection"  # 수동 선택으로
        else:
            # 불명확하거나 오류인 경우 다시 검증 요청
            return "process_verification_response"

    def manual_doc_type_router(self, state: State) -> str:
        """
        수동 문서 타입 선택 결과에 따라 다음 노드를 결정합니다.
        문서 타입이 이미 선택된 상태에서는 사용자 입력을 문서 내용으로 간주하고 check_user_input_policy로 이동합니다.
        
        Args:
            state (State): end_process, messages, doc_type 필드 포함
        
        Returns:
            str: "receive_user_input", "check_user_input_policy", "process_manual_doc_type_selection", "END" 중 하나
        """
        if state.get("end_process"):
            return "END"
        
        # 이미 문서 타입이 선택된 상태인지 확인
        if state.get("doc_type"):
            skip_ask_fields = state.get("skip_ask_fields", False)
            if skip_ask_fields:
                print(f"FLOW: 문서 타입 이미 선택됨 + 내용 있음: {state['doc_type']} - 바로 check_user_input_policy로 이동")
                return "check_user_input_policy"
            else:
                print(f"FLOW: 문서 타입 이미 선택됨: {state['doc_type']} - receive_user_input로 이동")
                return "receive_user_input"
        
        user_selection = ""
        if state.get("messages"):
            user_selection = state["messages"][-1].content.strip()
        
        # 문서 타입 매핑으로 유효성 검사
        doc_type_mapping = {
            "1": "영업방문 결과보고서",
            "2": "제품설명회 시행 신청서", 
            "3": "제품설명회 시행 결과보고서",
            "4": "종료",
            "영업방문 결과보고서": "영업방문 결과보고서",
            "제품설명회 시행 신청서": "제품설명회 시행 신청서",
            "제품설명회 시행 결과보고서": "제품설명회 시행 결과보고서",
            "종료": "종료"
        }
        
        selected_doc_type = doc_type_mapping.get(user_selection)
        skip_ask_fields = state.get("skip_ask_fields", False)
        
        if selected_doc_type == "종료":
            return "END"
        elif selected_doc_type:
            skip_ask_fields = state.get("skip_ask_fields", False)
            if skip_ask_fields:
                print("FLOW: 문서 타입 선택 완료 + 내용 있음 - 바로 check_user_input_policy로 이동")
                return "check_user_input_policy"
            else:
                print("FLOW: 문서 타입 선택 완료 - receive_user_input로 이동")
                return "receive_user_input"
        else:
            print(f"ERROR: 유효하지 않은 문서 타입 선택: '{user_selection}'")
            print("WARNING: 다음 중에서 정확히 선택해야 합니다:")
            print("  - 번호: 1, 2, 3, 4")
            print("  - 또는 정확한 문서명:")
            print("    * 영업방문 결과보고서")
            print("    * 제품설명회 시행 신청서")
            print("    * 제품설명회 시행 결과보고서")
            print("    * 종료")
            print("[INFO] 문서 타입 선택 실패로 인해 에이전트를 종료합니다.")
            return "END"


    def policy_check_router(self, state: State) -> str:
        """
        규정 검사 결과에 따라 다음 노드를 결정합니다.
        API 모드에서는 규정 위반이 있어도 분석을 계속 진행합니다.
        
        Args:
            state (State): violation 필드 포함
        
        Returns:
            str: "parse_user_input" 또는 "inform_violation"
        """
        violation = state.get("violation", "")
        
        if self.api_mode:
            # API 모드에서는 규정 위반 발견 시 즉시 종료
            if violation != "OK":
                print(f"[ERROR] 규정 위반이 발견되어 처리를 중단합니다 (API 모드)")
                return "inform_violation"
            else:
                print(f"[SUCCESS] 규정 위반 없음 - parse_user_input으로 이동")
                return "parse_user_input"
        else:
            # 콘솔 모드에서는 기존대로 동작
            if violation != "OK":
                print(f"[WARNING] 규정 위반 발견 - inform_violation으로 이동")
                return "inform_violation"
            else:
                print(f"[SUCCESS] 규정 위반 없음 - parse_user_input으로 이동")
                return "parse_user_input"
    
    def parse_router(self, state: State) -> str:
        """
        파싱 결과에 따라 다음 노드를 결정합니다.
        
        Args:
            state (State): parse_failed 필드 포함
        
        Returns:
            str: "check_user_input_policy" 또는 "create_choan_document"
        """
        if state.get("parse_failed"):
            return "check_user_input_policy"
        else:
            # 파싱 성공 시 문서 생성 단계로 이동
            violation = state.get("violation", "")
            
            # API 모드가 아니거나 규정 위반이 없는 경우에만 상세 로그 출력
            if not self.api_mode or violation == "OK":
                print("SUCCESS: 파싱 성공 - 문서 생성 단계로 이동")
                print("=" * 60)
                print("PARSED_DATA: 파싱된 사용자 입력 데이터:")
                print("=" * 60)
                
                filled_data = state.get("filled_data", {})
                for key, value in filled_data.items():
                    if value:
                        print(f"- {key}: {value}")
                
                print("=" * 60)
                print("SUCCESS: 문서 데이터 파싱 완료!")
            
            return "create_choan_document"
    


    def _build_graph(self):
        """
        LangGraph 워크플로우를 구성합니다.
        
        Returns:
            CompiledGraph: 컴파일된 LangGraph 워크플로우
        """
        graph = StateGraph(State)

        # 노드 추가
        graph.add_node("classify_doc_type", self.classify_doc_type)                          # 1️⃣ LLM으로 사용자 요청을 분석하여 문서 타입 분류
        graph.add_node("validate_doc_type", self.validate_doc_type)                          # 2️⃣ 분류된 문서 타입이 지원 문서인지 검증 (실패 시 수동 선택으로 이동)
        graph.add_node("process_verification_response", self.process_verification_response)  # 3️⃣ 사용자 검증 응답을 LLM으로 분석 (긍정/부정 판단)
        graph.add_node("process_manual_doc_type_selection", self.process_manual_doc_type_selection)  # 4️⃣ 사용자가 선택한 문서 타입 처리 및 템플릿 설정
        graph.add_node("receive_user_input", self.receive_user_input)                        # 5️⃣ 사용자 입력 수신
        graph.add_node("check_user_input_policy", self.check_user_input_policy)              # 6️⃣ 사용자 입력 텍스트로 규정 위반 검사 (LLM+OpenSearch)
        graph.add_node("parse_user_input", self.parse_user_input)                            # 7️⃣ 사용자 입력을 LLM으로 파싱하여 구조화된 JSON 데이터로 변환
        graph.add_node("inform_violation", self.inform_violation)                            # 8️⃣ 규정 위반 발견 시 위반 내용 안내 및 프로세스 종료
        graph.add_node("create_choan_document", self.create_choan_document)                  # 9️⃣ 파싱된 데이터로 DOCX 템플릿 기반 최종 문서 생성 및 저장

        # 흐름 연결
        graph.set_entry_point("classify_doc_type")
        
        # 분류 → 검증
        graph.add_edge("classify_doc_type", "validate_doc_type")
        
        # 검증 결과에 따른 분기
        graph.add_conditional_edges(
            "validate_doc_type",
            self.doc_type_validation_router,
            {
                "process_verification_response": "process_verification_response",  # 분류 성공 시 검증 단계로
                "process_manual_doc_type_selection": "process_manual_doc_type_selection"  # 분류 실패 시 수동 선택으로
            }
        )
        
        # 검증 응답 처리 결과에 따른 분기
        graph.add_conditional_edges(
            "process_verification_response",
            self.verification_response_router,
            {
                "receive_user_input": "receive_user_input",  # 긍정 + 내용 없음: 사용자 입력 수신
                "check_user_input_policy": "check_user_input_policy",  # 긍정 + 내용 있음: 바로 규정 검사
                "process_manual_doc_type_selection": "process_manual_doc_type_selection",  # 부정: 수동 선택
                "process_verification_response": "process_verification_response"  # 불명확: 다시 검증
            }
        )
        
        # 수동 선택 처리 결과에 따른 분기
        graph.add_conditional_edges(
            "process_manual_doc_type_selection",
            self.manual_doc_type_router,
            {
                "receive_user_input": "receive_user_input",  # 유효한 선택 + 내용 없음: 사용자 입력 수신
                "check_user_input_policy": "check_user_input_policy",  # 유효한 선택 + 내용 있음: 바로 규정 검사
                "process_manual_doc_type_selection": "process_manual_doc_type_selection",  # 유효하지 않은 선택
                "END": END  # 종료 선택
            }
        )
        
        # receive_user_input에서 check_user_input_policy로 연결
        graph.add_edge("receive_user_input", "check_user_input_policy")
        
        # 규정 검사 결과에 따른 분기
        graph.add_conditional_edges(
            "check_user_input_policy",
            self.policy_check_router,
            {
                "parse_user_input": "parse_user_input",
                "inform_violation": "inform_violation"
            }
        )
        
        # 파싱 결과에 따른 분기
        graph.add_conditional_edges(
            "parse_user_input",
            self.parse_router,
            {
                "check_user_input_policy": "check_user_input_policy",
                "create_choan_document": "create_choan_document"
            }
        )
        
        # 규정 위반 시 종료
        graph.add_edge("inform_violation", END)
        
        # 문서 생성 완료 후 종료
        graph.add_edge("create_choan_document", END)

        # 체크포인트 저장소 설정
        saver = MemorySaver()
        return graph.compile(
            checkpointer=saver, 
            interrupt_before=[
                "process_verification_response",      # 분류 검증용 인터럽트
                "process_manual_doc_type_selection",  # 수동 선택용 인터럽트
                "receive_user_input"                  # 사용자 입력 수신용 인터럽트
            ]
        )
    
    def run(self, user_input: str = None):
        """
        통합 문서 작성 시스템을 실행합니다.
        
        Args:
            user_input (str, optional): 사용자 입력. None이면 대화형 모드로 시작
        
        Returns:
            dict: 실행 결과 (success, result, thread_id, error 필드 포함)
        """
        
        # user_input이 없으면 대화형 모드로 시작
        if user_input is None:
            print("SYSTEM: 통합 문서 작성 시스템")
            print("=" * 60)
            print("SUPPORTED_DOCS: 지원 문서 타입:")
            print("  1. 영업방문 결과보고서")
            print("  2. 제품설명회 시행 신청서")
            print("  3. 제품설명회 시행 결과보고서")
            print("=" * 60)
            print("\n문서 작성 요청을 입력해주세요:")
            print(">>> ", end="")
            
            # 사용자 입력을 대기하고 받는 함수 호출
            user_input = self._get_user_input()
            
            print(f"\nPROCESSING_START: 처리 시작: {user_input}")
            print("=" * 60)
        
        # 초기 상태 설정
        initial_state = {
            "messages": [HumanMessage(content=user_input)],
            "doc_type": None,
            "template_content": None,
            "filled_data": None,
            "violation": None,
            "final_doc": None,
            "classification_failed": None,
            "skip_verification": None,
            "end_process": None,
            "parse_retry_count": None,
            "parse_failed": None,
            "user_reply": None,
            "verification_reply": None,
            "verification_result": None,
            "user_content": None,
            "skip_ask_fields": None
        }
        
        # 고유한 스레드 ID 생성
        thread_id = str(uuid.uuid4())
        config = {"configurable": {"thread_id": thread_id}}
        
        try:
            # 그래프 실행 (인터럽트 발생 시 중단)
            result = self.app.invoke(initial_state, config)
            
            # 최종 상태 확인
            violation_text = result.get("violation", "")
            has_no_violation = (violation_text == "OK")
            
            if has_no_violation and result.get("filled_data") and result.get("final_doc"):
                print("\n" + "="*50)
                print("DOCUMENT_COMPLETE: 문서 생성 완료!")
                print("="*50)
                
                result_json = json.dumps(result["filled_data"], indent=2, ensure_ascii=False)
                print(result_json)
                print(f"\nFINAL_DOC: 생성된 문서: {result['final_doc']}")
                
                return {"success": True, "result": result, "thread_id": thread_id}
            else:
                # 인터럽트로 중단된 경우
                print(f"\nINTERRUPT: 인터럽트 발생 - 스레드 ID: {thread_id}")
                if self.api_mode:
                    # API 모드일 경우 인터럽트 정보 반환
                    return self._handle_api_interrupt(thread_id)
                else:
                    # 콘솔 모드일 경우 기존 대화형 처리
                    return self._handle_interactive_mode(thread_id)
                
        except Exception as e:
            print(f"\nERROR: 실행 중 오류: {e}")
            return {"success": False, "error": str(e)}
    
    def _get_user_input(self) -> str:
        """
        사용자 입력을 받는 통합 메서드
        
        Returns:
            str: 사용자 입력 문자열
        """
        try:
            return input()
        except KeyboardInterrupt:
            print("\n\nUSER_CANCELLED: 사용자가 중단했습니다.")
            raise
        except Exception as e:
            print(f"\nERROR: 입력 오류: {e}")
            raise
    
    def _handle_interactive_mode(self, thread_id: str):
        """
        인터럽트 발생 시 대화형 모드 처리
        
        Args:
            thread_id (str): 스레드 ID
        
        Returns:
            dict: 처리 결과 (success, result, interrupted_by_user, error 필드 포함)
        """
        print(f"INTERRUPT_HANDLER: 인터럽트 발생 - 스레드 ID: {thread_id}")
        
        # 인터럽트 처리 루프
        while True:
            try:
                # 현재 상태 확인하여 어떤 노드에서 인터럽트가 발생했는지 확인
                current_state = self.app.get_state({"configurable": {"thread_id": thread_id}})
                next_node = current_state.next[0] if current_state.next else None
                
                # 다음 노드에 따라 안내 문구 출력 및 입력 타입 결정
                if next_node == "process_verification_response":
                    print("\nCLASSIFICATION_RESULT: 문서 타입 분류 결과 확인")
                    print("=" * 60)
                    doc_type = current_state.values.get("doc_type", "")
                    print(f"DOC_TYPE: 분류된 문서 타입: {doc_type}")
                    print("=" * 60)
                    print("\n위 분류 결과가 올바른가요?")
                    print("- 맞다면 'YES' 또는 '네' 또는 '맞습니다' 등으로 응답해주세요")
                    print("- 틀렸다면 'NO' 또는 '아니요' 또는 '틀렸습니다' 등으로 응답해주세요")
                    input_type = "verification_reply"
                elif next_node == "process_manual_doc_type_selection":
                    print("\nMANUAL_SELECT: 올바른 문서 타입을 선택해주세요:")
                    print("=" * 60)
                    print("1. 영업방문 결과보고서")
                    print("2. 제품설명회 시행 신청서") 
                    print("3. 제품설명회 시행 결과보고서")
                    print("4. 종료")
                    print("=" * 60)
                    print("\n위 번호(1-4) 또는 문서명을 직접 입력해주세요.")
                    input_type = "verification_reply"  # 수동 선택도 verification_reply 사용
                elif next_node == "receive_user_input":
                    # 문서 작성을 위한 필수 항목 안내
                    template_content = current_state.values.get("template_content", "")
                    if template_content:
                        print("\nREQUIRED_FIELDS: 다음 항목들을 입력해주세요:")
                        print("=" * 60)
                        print(template_content)
                        print("=" * 60)
                        print("\n위 항목들에 맞춰 정보를 입력해주세요.")
                    else:
                        print("\nINPUT_REQUEST: 문서 작성을 위한 정보를 입력해주세요.")
                    input_type = "user_reply"
                else:
                    print("\nINPUT_REQUEST: 문서 작성을 위한 정보를 입력해주세요.")
                    input_type = "user_reply"  # 기본값
                
                print("WAITING_INPUT: 사용자 입력을 기다립니다.")
                print(">>> ", end="")
                user_response = self._get_user_input()
                
                # 에이전트 재개
                resume_result = self.resume(thread_id, user_response, input_type)
                
                if resume_result.get("success"):
                    print("\nCOMPLETE: 문서 작성 완료!")
                    return {"success": True, "result": resume_result.get("result")}
                elif resume_result.get("interrupted"):
                    # 또 다른 인터럽트가 발생한 경우 계속 진행
                    continue
                else:
                    print(f"\n❌ 처리 실패: {resume_result}")
                    return {"success": False, "result": resume_result}
                    
            except KeyboardInterrupt:
                print("\n\nUSER_CANCELLED: 사용자가 중단했습니다.")
                return {"success": False, "interrupted_by_user": True}
            except Exception as e:
                print(f"\n❌ 오류 발생: {e}")
                return {"success": False, "error": str(e)}
    
    def resume(self, thread_id: str, user_reply: str, input_type: str = "user_reply"):
        """
        인터럽트된 워크플로우를 사용자 입력과 함께 재개합니다.
        
        Args:
            thread_id (str): 스레드 ID
            user_reply (str): 사용자 입력
            input_type (str): 입력 타입 ("user_reply", "verification_reply")
        
        Returns:
            dict: 재개 결과 (success, result, interrupted, thread_id, error 필드 포함)
        """
        config = {"configurable": {"thread_id": thread_id}}
        
        try:
            # 현재 상태 가져오기
            current_state = self.app.get_state(config)
            print(f"CURRENT_STATE: 현재 상태: {current_state}")
            
            # 사용자 입력을 상태에 업데이트 (입력 타입에 따라)
            update_data = {input_type: user_reply}
            self.app.update_state(config, update_data)
            
            # 사용자 입력을 메시지 히스토리에도 추가
            new_message = HumanMessage(content=user_reply)
            current_messages = current_state.values.get("messages", [])
            current_messages.append(new_message)
            self.app.update_state(config, {"messages": current_messages})
            
            print(f"INPUT_TYPE: {input_type}, 사용자 입력: {user_reply[:50]}...")
            
            # 워크플로우 재개 - API 모드 체크
            if self.api_mode:
                # API 모드에서는 invoke 사용하여 한 번에 처리
                try:
                    final_result = self.app.invoke(None, config)
                    
                    # 디버깅을 위해 final_result 내용 출력
                    print(f"[DEBUG] final_result keys: {final_result.keys() if final_result else 'None'}")
                    print(f"[DEBUG] final_result.violation: {final_result.get('violation', 'NOT FOUND')}")
                    print(f"[DEBUG] final_result.error_type: {final_result.get('error_type', 'NOT FOUND')}")
                    
                    # 또 다른 인터럽트 확인
                    current_state_after = self.app.get_state(config)
                    if current_state_after.next:
                        return self._handle_api_interrupt(thread_id)
                    
                    # 종료 처리 확인
                    if final_result.get("end_process"):
                        # 위반사항 먼저 확인
                        violation_text = final_result.get("violation", "")
                        error_type = final_result.get("error_type", "")
                        
                        # 중간 종료 시 그래프 상태를 완전히 정리
                        print(f"[INFO] 중간 종료 감지, 그래프 상태 정리: {thread_id}")
                        try:
                            # 그래프를 강제로 END 상태로 만들어 완료 처리
                            self.app.update_state(config, {"end_process": True})
                            # 현재 상태를 END로 마킹하여 이후 resume 시 새로운 대화로 인식되도록 함
                            print(f"[INFO] 그래프 상태 정리 완료: {thread_id}")
                        except Exception as e:
                            print(f"[ERROR] 그래프 상태 정리 중 오류: {e}")
                        
                        # 위반사항이 있는 경우
                        if violation_text and violation_text != "OK":
                            return {
                                "success": False,
                                "result": final_result,
                                "thread_id": thread_id,
                                "error_type": "policy_violation",
                                "violation": violation_text,
                                "message": "규정 위반으로 문서 생성이 중단되었습니다.",
                                "end_process": True
                            }
                        # 규정 위반 타입으로 명시된 경우
                        elif error_type == "policy_violation":
                            return {
                                "success": False,
                                "result": final_result,
                                "thread_id": thread_id,
                                "error_type": "policy_violation",
                                "violation": violation_text or "규정 위반",
                                "message": "규정 위반으로 문서 생성이 중단되었습니다.",
                                "end_process": True
                            }
                        # 사용자 종료
                        else:
                            return {
                                "success": False,
                                "result": final_result,
                                "thread_id": thread_id,
                                "error_type": "user_terminated",
                                "message": "사용자가 종료를 선택했습니다.",
                                "end_process": True
                            }
                    
                    # 성공 또는 실패 처리
                    violation_text = final_result.get("violation", "")
                    has_violation = (violation_text != "OK")
                    
                    # API 모드에서는 규정 위반이 있어도 분석이 완료되면 성공으로 처리
                    if final_result.get("filled_data"):
                        if has_violation:
                            # 규정 위반이 있지만 분석은 완료됨
                            return {
                                "success": True,  # 분석은 성공했으므로 True
                                "result": final_result,
                                "thread_id": thread_id,
                                "filled_data": final_result.get("filled_data"),
                                "final_doc": final_result.get("final_doc"),  # 파일 생성이 차단되면 None
                                "violation": violation_text,
                                "violation_blocked": final_result.get("violation_blocked", False)
                            }
                        else:
                            # 규정 위반 없이 성공
                            return {
                                "success": True, 
                                "result": final_result,
                                "thread_id": thread_id,
                                "filled_data": final_result.get("filled_data"),
                                "final_doc": final_result.get("final_doc")
                            }
                    else:
                        # 분석 자체가 실패한 경우
                        return {
                            "success": False, 
                            "result": final_result,
                            "thread_id": thread_id,
                            "error_type": "processing_error"
                        }
                        
                except Exception as e:
                    return {"success": False, "error": str(e), "thread_id": thread_id}
            else:
                # 기존 콘솔 모드 처리 - stream 사용
                final_result = None
                for chunk in self.app.stream(None, config):
                    print(f"PROCESSING: 처리 중: {list(chunk.keys())}")
                    if chunk:
                        final_result = list(chunk.values())[-1]  # 마지막 결과 저장
                
                # 최종 상태 확인
                if final_result:
                    violation_text = final_result.get("violation", "")
                    has_no_violation = (violation_text == "OK")
                    
                    if has_no_violation and final_result.get("filled_data") and final_result.get("final_doc"):
                        print("\n" + "="*50)
                        print("DOCUMENT_COMPLETE: 문서 생성 완료!")
                        print("="*50)
                        
                        result_json = json.dumps(final_result["filled_data"], indent=2, ensure_ascii=False)
                        print(result_json)
                        print(f"\n📁 생성된 문서: {final_result['final_doc']}")
                        
                        return {"success": True, "result": final_result}
                else:
                    # 중간 인터럽트 상황도 처리
                    current_state_after = self.app.get_state(config)
                    if current_state_after.next:  # 다음 실행할 노드가 있으면 인터럽트 상황
                        next_node = current_state_after.next[0] if current_state_after.next else None
                        print(f"NEXT_INTERRUPT: 다음 인터럽트 대기 중 - 다음 노드: {next_node}")
                        return {"success": False, "interrupted": True, "thread_id": thread_id, "next_node": next_node}
                    else:
                        print("\n❌ 문서 생성 실패")
                        print(f"최종 결과: {final_result}")
                        return {"success": False, "result": final_result}
                
        except Exception as e:
            print(f"\n❌ 재개 중 오류: {e}")
            import traceback
            traceback.print_exc()
            return {"success": False, "error": str(e)}
    
    def _handle_api_interrupt(self, thread_id: str):
        """
        API 모드에서 인터럽트 처리
        
        Args:
            thread_id (str): 스레드 ID
        
        Returns:
            dict: 인터럽트 정보 (interrupted, thread_id, next_node, doc_type 등 포함)
        """
        current_state = self.app.get_state({"configurable": {"thread_id": thread_id}})
        next_node = current_state.next[0] if current_state.next else None
        
        # 현재 상태에서 모든 정보 추출
        state_values = current_state.values
        doc_type = state_values.get("doc_type")
        
        # 전체 state 정보를 포함하여 누락 방지
        interrupt_info = {
            "success": False,
            "interrupted": True,
            "thread_id": thread_id,
            "next_node": next_node,
            "doc_type": doc_type,
            "state_info": {
                # 모든 state 필드 포함
                "messages": [msg.content if hasattr(msg, 'content') else str(msg) for msg in state_values.get("messages", [])],
                "template_content": state_values.get("template_content"),
                "filled_data": state_values.get("filled_data", {}),
                "violation": state_values.get("violation"),
                "final_doc": state_values.get("final_doc"),
                "classification_failed": state_values.get("classification_failed"),
                "skip_verification": state_values.get("skip_verification"),
                "end_process": state_values.get("end_process"),
                "parse_retry_count": state_values.get("parse_retry_count"),
                "parse_failed": state_values.get("parse_failed"),
                "user_reply": state_values.get("user_reply"),
                "verification_reply": state_values.get("verification_reply"),
                "verification_result": state_values.get("verification_result"),
                "user_content": state_values.get("user_content"),
                "skip_ask_fields": state_values.get("skip_ask_fields")
            }
        }
        
        # 노드별 프롬프트 설정
        if next_node == "receive_verification_input":
            interrupt_info["prompt"] = f"분류된 문서 타입: {doc_type}\n\n위 분류 결과가 올바른가요?"
            interrupt_info["prompt_type"] = "verification"
        elif next_node == "receive_manual_doc_type_input":
            interrupt_info["prompt"] = "문서 타입을 선택해주세요."
            interrupt_info["prompt_type"] = "manual_selection"
            interrupt_info["options"] = [
                {"value": "1", "label": "영업방문 결과보고서"},
                {"value": "2", "label": "제품설명회 시행 신청서"},
                {"value": "3", "label": "제품설명회 시행 결과보고서"},
                {"value": "4", "label": "종료"}
            ]
        elif next_node == "receive_user_input":
            # 필요한 필드 정보 포함
            template_info = self.doc_prompts.get(doc_type, {})
            interrupt_info["prompt"] = template_info.get("input_prompt", "필요한 정보를 입력해주세요.")
            interrupt_info["prompt_type"] = "field_input"
            interrupt_info["required_fields"] = template_info.get("required_fields", [])
            # 템플릿 정보도 포함
            interrupt_info["template_content"] = state_values.get("template_content")
        
        return interrupt_info

    def get_current_state(self, thread_id: str):
        """특정 thread의 현재 상태를 조회합니다."""
        try:
            if hasattr(self, 'graph'):
                config = {"configurable": {"thread_id": thread_id}}
                state_snapshot = self.graph.get_state(config)
                return state_snapshot
            else:
                print(f"[WARNING] Graph not initialized for thread {thread_id}")
                return None
        except Exception as e:
            print(f"[ERROR] Error getting current state for thread {thread_id}: {e}")
            return None

    def clear_thread_state(self, thread_id: str) -> bool:
        """특정 thread의 상태를 초기화합니다."""
        try:
            # app 객체에서 checkpointer 접근
            if hasattr(self, 'app') and hasattr(self.app, 'checkpointer'):
                config = {"configurable": {"thread_id": thread_id}}
                
                # 그래프를 강제로 완료 상태로 만들기
                try:
                    current_state = self.app.get_state(config)
                    if current_state and current_state.values:
                        # 강제로 end_process를 True로 설정하고 그래프 완료
                        self.app.update_state(config, {"end_process": True})
                        print(f"[INFO] 그래프를 강제 완료 상태로 설정: {thread_id}")
                except Exception as e:
                    print(f"[WARNING] 그래프 강제 완료 설정 실패: {e}")
                
                # MemorySaver의 경우 직접 삭제
                if hasattr(self.app.checkpointer, 'storage'):
                    # 메모리 저장소에서 해당 thread_id 관련 데이터 삭제
                    storage = self.app.checkpointer.storage
                    keys_to_delete = []
                    for key in storage.keys():
                        if thread_id in str(key):
                            keys_to_delete.append(key)
                    
                    for key in keys_to_delete:
                        del storage[key]
                    
                    print(f"[INFO] Successfully cleared thread state: {thread_id} (deleted {len(keys_to_delete)} keys)")
                    return True
                else:
                    print(f"[WARNING] Checkpointer storage not accessible for thread {thread_id}")
                    return False
            else:
                print(f"[WARNING] App or checkpointer not available for thread {thread_id}")
                return False
                
        except Exception as e:
            print(f"[ERROR] Error clearing thread state {thread_id}: {e}")
            return False

if __name__ == "__main__":
    # 통합 문서 작성 시스템 실행
    agent = CreateDocumentAgent()
    agent.run(user_input="영업방문결과보고서를 작성해줘")