from langgraph.graph import StateGraph, END
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, ToolMessage, AIMessage
from langchain_core.tools import tool
from langchain_core.prompts import ChatPromptTemplate
from typing import Annotated, TypedDict, List, Optional
import os
import docx
import json
import yaml
import time
from pathlib import Path
from dotenv import load_dotenv
from docx import Document

load_dotenv()

class State(TypedDict):
    messages: List[HumanMessage]
    doc_type: Optional[str]
    template_content: Optional[str]
    filled_data: Optional[dict]
    violation: Optional[str]
    final_doc: Optional[str]
    retry_count: int
    restart_classification: Optional[bool]
    classification_retry_count: Optional[int]
    end_process: Optional[bool]
    parse_retry_count: Optional[int]
    parse_failed: Optional[bool]

class DocumentDraftAgent:
    """지능형 문서 초안 작성 시스템"""
    
    def __init__(self, model_name: str = "gpt-4o-mini", temperature: float = 0.7, document_model: str = "gpt-4o"):
        """
        DocumentDraftAgent 초기화
        
        Args:
            model_name: 기본 LLM 모델명 (파싱용)
            temperature: LLM 온도 설정
            document_model: 문서 생성용 모델명
        """
        self.model_name = model_name
        self.document_model = document_model
        self.temperature = temperature
        
        # 툴 정의
        self.tools = [self.check_policy_violation]
        
        # LLM 초기화 (파싱용)
        self.llm = ChatOpenAI(
            model=self.model_name, 
            temperature=self.temperature
        ).bind_tools(self.tools)
        
        # 문서 생성용 LLM 초기화
        self.document_llm = ChatOpenAI(
            model=self.document_model,
            temperature=0.0
        )
        
        # YAML 파일에서 템플릿 로드
        self.doc_prompts = self._load_templates()
        
        # 그래프 초기화
        self.app = self._build_graph()
    
    def _load_templates(self):
        """YAML 파일에서 템플릿을 로드합니다."""
        try:
            current_dir = Path(__file__).parent
            template_path = current_dir / "templates.yaml"
            
            if not template_path.exists():
                print(f"⚠️ 템플릿 파일을 찾을 수 없습니다: {template_path}")
                return {}
            
            with open(template_path, 'r', encoding='utf-8') as file:
                data = yaml.safe_load(file)
                return data.get('templates', {})
                
        except Exception as e:
            print(f"❌ 템플릿 로드 중 오류 발생: {e}")
            return {}
        
    
    @staticmethod
    @tool
    def check_policy_violation(content: Annotated[str, "작성된 문서 본문"]) -> str:
        """작성된 문서 내용이 회사 규정을 위반하는지 검사합니다."""
        # 실제 규정 검사 로직 (예시)
        violations = []
        
        # 금지어 체크
        forbidden_words = ["금지어", "부정적", "비밀", "기밀유출"]
        for word in forbidden_words:
            if word in content:
                violations.append(f"금지어 포함: '{word}'")
        
        # 기본적인 필수 항목만 체크 (너무 엄격하지 않게)
        if len(content.strip()) < 10:
            violations.append("입력 내용이 너무 짧습니다")
        
        # 최소한의 정보 확인
        basic_info_found = any(keyword in content for keyword in ["방문", "고객", "회사", "협의", "논의", "만나", "시행", "제품"])
        if not basic_info_found:
            violations.append("방문 관련 기본 정보가 부족합니다")
        
        if violations:
            return " | ".join(violations)
        return "OK"


    def parse_user_input(self, state: State) -> State:
        user_input = str(state["messages"][-1].content)
        doc_type = state["doc_type"]
        response = None  # 에러 발생 대비 초기화

        if state.get("parse_retry_count") is None:
            state["parse_retry_count"] = 0

        system_prompt = self.doc_prompts[doc_type]["choan_system_prompt"]
        if not system_prompt:
            raise ValueError(f"문서 타입에 대한 시스템 프롬프트가 없습니다: {doc_type}")

        # 중괄호 이스케이프 처리
        escaped_input = user_input.replace("{", "{{").replace("}", "}}")

        parsing_prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{user_input}")
        ])

        try:
            formatted_messages = parsing_prompt.format_messages(user_input=escaped_input)
            print("📨 LLM에 전달된 메시지:")
            for m in formatted_messages:
                print(f"[{m.type.upper()}] {m.content[:200]}...")

            response = self.llm.invoke(formatted_messages)

            content = response.content
            json_str = content if isinstance(content, str) else str(content)
            print(f"\n🔍 LLM 응답 내용:\n{json_str}")

            if "{" in json_str and "}" in json_str:
                start = json_str.find("{")
                end = json_str.rfind("}") + 1
                clean_json = json_str[start:end]
                print(f"\n🔍 추출된 JSON:\n{clean_json}")

                import json
                try:
                    parsed_data = json.loads(clean_json)
                    state["filled_data"] = parsed_data
                    state["parse_failed"] = False
                    print("✅ 파싱 성공:", parsed_data)
                except json.JSONDecodeError as json_error:
                    print(f"❌ JSON 파싱 오류: {json_error}")
                    print(f"파싱 시도한 JSON: {repr(clean_json)}")
                    raise json_error
            else:
                raise ValueError("구조화된 JSON 형식을 찾을 수 없음")

        except Exception as e:
            print("\n⚠️ 예외 발생!")
            if response:
                print("응답 내용:")
                print(response)
            else:
                print("⚠️ response 객체가 존재하지 않습니다.")
            print(f"⚠️ 예외 메시지: {e}")

            retry_count = state.get("parse_retry_count", 0) + 1
            state["parse_retry_count"] = retry_count

            if retry_count >= 3:
                print("⚠️ 파싱 재시도 초과. 기본값 사용.")
                fallback_data = self.doc_prompts[doc_type]["choan_fallback_fields"]
                state["filled_data"] = fallback_data
            else:
                print(f"🔄 재시도 {retry_count}/3")
                state["parse_failed"] = True

        return state

    def run_check_policy_violation(self, state: State) -> State:
        """작도 데이터가 규정을 위반하는지 검사합니다."""
        filled_data = state["filled_data"] or {}
        content = " ".join(str(v) for v in filled_data.values())
        
        try:
            result = self.check_policy_violation.invoke({"content": content})
            state["violation"] = result
            print(f"🔍 규정 검사 결과: {result}")
            
            # 규정 위반이 없으면 parse_user_input 결과를 출력
            if result == "OK":
                print("\n✅ 규정 위반이 없습니다!")
                print("=" * 60)
                print("📝 파싱된 사용자 입력 데이터:")
                print("=" * 60)
                
                for key, value in filled_data.items():
                    if value:  # 빈 값이 아닌 경우만 출력
                        print(f"- {key}: {value}")
                
                print("=" * 60)
                print("✅ 문서 데이터 파싱 완료!")
                return state
            else:
                print("❌ 규정 위반이 발견되었으므로 재입력을 요청합니다.")
                return state
        
        except Exception as e:
            print(f"⚠️ 규정 검사 실패: {e}")
            state["violation"] = "규정 검사 중 오류가 발생했습니다. 내용을 다시 입력해주세요."
            return state

    def inform_violation(self, state: State) -> State:
        """규정 위반이 발견되었을 때 자동으로 재시도합니다."""
        violation = state["violation"]
        retry_count = state.get("retry_count", 0) + 1
        state["retry_count"] = retry_count
        
        print(f"\n⚠️ 규정 위반 사항 발견 (시도 #{retry_count}/3):")
        print(f"문제점: {violation}")
        print("자동으로 재시도합니다...")
        
        # 자동 재시도 - 원본 사용자 입력 유지
        original_input = state["messages"][-1].content if state["messages"] else ""
        if original_input:
            print(f"원본 입력으로 재시도: {original_input[:100]}...")
        
        return state

    def create_choan_document(self, state: State) -> State:
        """파싱된 데이터를 기반으로 초안 문서를 생성하고 docx 파일로 저장합니다."""
        doc_type = state["doc_type"]
        filled_data = state["filled_data"]
        
        # 문서 타입에 따른 템플릿 파일 매핑
        template_mapping = {
            "영업방문 결과보고서": "영업방문 결과보고서(템플릿형).docx",
            "제품설명회 시행 신청서": "제품설명회 시행 신청서(템플릿형).docx",
            "제품설명회 시행 결과보고서": "제품설명회 시행 결과보고서(템플릿형).docx"
        }
        template_filename = template_mapping.get(doc_type)
        if not template_filename:
            print(f"❌ 지원하지 않는 문서 타입: {doc_type}")
            state["final_doc"] = None
            return state
        
        # S3 폴더에서 템플릿 파일 경로 구성
        current_dir = Path(__file__).parent
        template_path = current_dir / "S3" / template_filename
        
        if not template_path.exists():
            print(f"❌ 템플릿 파일을 찾을 수 없습니다: {template_path}")
            state["final_doc"] = None
            return state
        
        try:
            # 템플릿 파일 읽기
            print(f"📂 템플릿 파일 로딩: {template_filename}")
            doc = Document(str(template_path))
            
            print(f"📝 템플릿 플레이스홀더 치환 중...")
            
            # 양식을 유지하면서 플레이스홀더만 치환
            self._replace_placeholders_in_document(doc, filled_data, doc_type)
            
            # agent_result_document_folder 디렉토리 생성 (없으면 생성)
            result_folder = current_dir / "agent_result_document_folder"
            result_folder.mkdir(exist_ok=True)
            
            # 완성된 문서 저장
            today_date = time.strftime('%Y%m%d')
            # 문서 타입에서 띄어쓰기 제거
            doc_type_no_space = doc_type.replace(" ", "")
            output_filename = f"{doc_type_no_space}_{today_date}.docx"
            output_path = result_folder / output_filename
            doc.save(str(output_path))
            
            state["final_doc"] = str(output_path)
            
            print("✅ 문서 생성 및 저장 완료!")
            print(f"📁 저장 경로: {output_path}")
            print("📝 템플릿 양식이 그대로 유지되면서 플레이스홀더만 치환되었습니다.")
            
        except Exception as e:
            print(f"❌ 문서 생성 중 오류 발생: {e}")
            import traceback
            traceback.print_exc()
            state["final_doc"] = None
        
        return state

    def _replace_placeholders_in_document(self, doc, filled_data, doc_type):
        """문서의 플레이스홀더를 실제 데이터로 치환합니다."""
        
        # 개별 예산 플레이스홀더로 변경됨에 따라 특별 처리 불필요
        
        # 다중 항목 처리를 위한 특별 처리 필요 항목들 (템플릿과 데이터 매칭)
        # 문서 타입에 따라 다른 플레이스홀더 매핑 사용
        if doc_type == "제품설명회 시행 신청서":
            multi_item_fields = {
                "직원팀명": "직원팀명", 
                "팀명성명": "직원성명",  # 신청서에서는 팀명성명항목내용 사용
                "의료기관명": "의료기관명",
                "보건의료전문가성명": "보건의료전문가성명"
            }
        else:
            multi_item_fields = {
                "참석직원팀명": "직원팀명", 
                "참석직원성명": "직원성명",
                "참석의료기관명": "의료기관명",
                "참석보건의료전문가성명": "보건의료전문가성명"
            }
        
        # 문서에서 실제로 사용되는 플레이스홀더 번호 범위를 동적으로 찾기
        max_placeholders = self._find_max_placeholder_numbers(doc, multi_item_fields.keys())
        
        # 일반 플레이스홀더 치환 (문단)
        for paragraph in doc.paragraphs:
            self._replace_in_text_element(paragraph, filled_data, multi_item_fields, max_placeholders)
        
        # 테이블 내용 치환
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_text_element(paragraph, filled_data, multi_item_fields, max_placeholders)

    def _find_max_placeholder_numbers(self, doc, field_keys):
        """문서에서 실제로 사용되는 최대 플레이스홀더 번호를 찾습니다."""
        import re
        max_numbers = {}
        
        # 모든 텍스트 수집
        all_text = ""
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_text += paragraph.text + "\n"
        
        # 각 필드별 최대 번호 찾기
        for field_key in field_keys:
            pattern = rf"{field_key}항목내용(\d+)"
            numbers = re.findall(pattern, all_text)
            if numbers:
                max_numbers[field_key] = max(int(n) for n in numbers)
            else:
                max_numbers[field_key] = 0
        
        return max_numbers

    def _replace_in_text_element(self, text_element, filled_data, multi_item_fields, max_placeholders):
        """텍스트 요소에서 플레이스홀더를 치환합니다. 포맷팅을 유지하면서 치환합니다."""
        
        # 모든 치환 작업을 수집
        replacements = {}
        
        # 일반 필드 처리  
        for key, value in filled_data.items():
            if key not in multi_item_fields.values():
                # 지급내역은 특별 처리 (플레이스홀더가 제품설명회지급내역항목내용)
                if key == "지급내역":
                    placeholder = "제품설명회지급내역항목내용"
                    replacement_value = str(value) if value else ""
                    replacements[placeholder] = replacement_value
                # 개별 예산 필드들 처리 (1인금액을 먼저 처리하여 겹침 방지)
                elif key in ["1인금액", "금액", "메뉴", "주류"]:
                    placeholder = f"{key}항목내용"
                    replacement_value = str(value) if value else ""
                    replacements[placeholder] = replacement_value
                else:
                    placeholder = f"{key}항목내용"
                    replacement_value = str(value) if value else ""
                    replacements[placeholder] = replacement_value
        
        # 다중 항목 필드 처리
        for field_key, data_key in multi_item_fields.items():
            value = filled_data.get(data_key, "")  # 데이터가 없으면 빈 문자열
            # 콤마로 분리하여 리스트로 변환
            items = [item.strip() for item in str(value).split(',')] if value else []
            
            # 동적으로 찾은 최대 번호까지 처리
            max_num = max_placeholders.get(field_key, 0)
            for i in range(1, max_num + 1):
                placeholder = f"{field_key}항목내용{i}"
                replacement_value = items[i-1] if i-1 < len(items) else ""
                replacements[placeholder] = replacement_value
        
        # 템플릿에 있는 추가 플레이스홀더들 처리 (번호 없는 것들)
        additional_placeholders = [
            "PM참석항목내용", "구분항목내용", "일시항목내용", "장소항목내용", 
            "제품명항목내용", "제품설명회시행목적항목내용", "제품설명회주요내용항목내용", 
            "참석인원항목내용", "방문일항목내용"
        ]
        
        for placeholder in additional_placeholders:
            if placeholder not in replacements:
                # 해당하는 데이터 키 찾기
                data_key = placeholder.replace("항목내용", "")
                # 특별한 매핑 처리
                if placeholder == "방문일항목내용":
                    data_key = "방문날짜"
                replacement_value = str(filled_data.get(data_key, ""))
                replacements[placeholder] = replacement_value
        
        # run 단위로 포맷팅을 유지하면서 치환
        self._replace_text_preserving_format(text_element, replacements)

    def _replace_text_preserving_format(self, paragraph, replacements):
        """포맷팅을 유지하면서 텍스트를 치환합니다."""
        if not replacements:
            return
            
        # 모든 run에서 텍스트를 수집
        full_text = ""
        run_texts = []
        
        for run in paragraph.runs:
            run_text = run.text
            run_texts.append(run_text)
            full_text += run_text
        
        # 치환 작업 수행 - 정확한 매칭으로 겹침 방지
        modified_text = full_text
        for placeholder, replacement in replacements.items():
            if placeholder in modified_text:
                # 특별 처리: 금액항목내용이 1인금액항목내용의 일부인지 확인
                if placeholder == "금액항목내용":
                    # 1인금액항목내용이 아닌 독립적인 금액항목내용만 치환
                    if "1인금액항목내용" not in modified_text:
                        modified_text = modified_text.replace(placeholder, replacement)
                    else:
                        # 1인금액항목내용을 제외한 나머지 금액항목내용만 치환
                        import re
                        # 1인금액항목내용이 아닌 금액항목내용만 매칭하는 패턴
                        pattern = r'(?<!1인)금액항목내용'
                        modified_text = re.sub(pattern, replacement, modified_text)
                else:
                    modified_text = modified_text.replace(placeholder, replacement)
        
        # 텍스트가 변경되었을 때만 처리
        if modified_text != full_text:
            # 모든 기존 run 제거
            for run in paragraph.runs[:]:
                run._element.getparent().remove(run._element)
            
            # 새로운 run으로 변경된 텍스트 추가 (첫 번째 run의 포맷팅 유지)
            if paragraph.runs or run_texts:
                # 첫 번째 run이 있었다면 그 포맷팅을 기본으로 사용
                new_run = paragraph.add_run(modified_text)
                # 기존 스타일 복사는 복잡하므로 기본 동작에 맡김
            else:
                # run이 없었다면 단순히 텍스트만 추가
                paragraph.add_run(modified_text)

    def doc_type_validation_router(self, state: State) -> str:
        """문서 타입 유효성 검사 결과에 따라 다음 노드를 결정합니다."""
        doc_type = state.get("doc_type", "")
        valid_types = ["영업방문 결과보고서", "제품설명회 시행 신청서", "제품설명회 시행 결과보고서"]
        retry_count = state.get("classification_retry_count") or 0
        
        if state.get("end_process"):
            return "END"
        elif doc_type in valid_types:
            return "ask_required_fields"
        else:
            return "classify_doc_type"
    
    def parse_router(self, state: State) -> str:
        """파싱 결과에 따라 다음 노드를 결정합니다."""
        if state.get("parse_failed"):
            return "ask_required_fields"
        else:
            return "check_policy_violation"

    def policy_check_router(self, state: State) -> str:
        """규정 검사 결과에 따라 다음 노드를 결정합니다."""
        if state.get("violation") == "OK":
            return "create_choan_document"
        else:
            # 재시도 횟수 제한 (무한 루프 방지)
            retry_count = state.get("retry_count", 0)
            if retry_count >= 3:
                print("⚠️ 최대 재시도 횟수(3회) 초과, 처리를 종료합니다.")
                return "END"
            return "inform_violation"

    def _build_graph(self):
        """LangGraph 워크플로우를 구성합니다."""
        graph = StateGraph(State)

        # 노드 추가 (ask_required_fields 제거)
        graph.add_node("parse_user_input", self.parse_user_input)
        graph.add_node("check_policy_violation", self.run_check_policy_violation)
        graph.add_node("inform_violation", self.inform_violation)
        graph.add_node("create_choan_document", self.create_choan_document)

        # 흐름 연결 - parse_user_input부터 시작
        graph.set_entry_point("parse_user_input")

        # 파싱 결과에 따른 분기 (파싱 실패시 재시도는 내부적으로 처리)
        graph.add_edge("parse_user_input", "check_policy_violation")

        # 조건부 분기 - 규정 위반 시 재입력 루프, OK시 초안 문서 생성
        graph.add_conditional_edges(
            "check_policy_violation",
            self.policy_check_router,
            {
                "END": END,
                "inform_violation": "inform_violation",
                "create_choan_document": "create_choan_document"
            }
        )

        # 규정 위반 시 재입력 → 파싱 → 검사 루프
        graph.add_edge("inform_violation", "parse_user_input")
        
        # 초안 문서 생성 완료 후 종료
        graph.add_edge("create_choan_document", END)

        return graph.compile()
    
    def run(self):
        """워크플로우를 실행하고 결과를 반환합니다."""
        initial_state = {
            "messages": [],
            "doc_type": '영업방문 결과보고서',
            "template_content": None,
            "filled_data": None,
            "violation": None,
            "final_doc": None,
            "retry_count": 0,
            "restart_classification": None,
            "classification_retry_count": None
        }
        
        # 그래프 실행
        final_state = self.app.invoke(initial_state)
        
        # 최종 상태가 성공적인지 확인 (규정 위반 없음 + 데이터 존재)
        if final_state.get("violation") == "OK" and final_state.get("filled_data"):
            print("\n" + "="*50)
            print("📄 최종 파싱 결과:")
            print("="*50)
            
            import json
            result = json.dumps(final_state["filled_data"], indent=2, ensure_ascii=False)
            print(result)
            
            return final_state["filled_data"]
        else:
            # 실패 메시지는 각 노드/라우터에서 이미 출력되었으므로, 여기서는 None만 반환하여
            # __main__ 블록에서 "처리 실패"가 출력되도록 함.
            return None
    
    def run_with_state(self, input_state: dict, user_input: str):
        """기존 state와 사용자 입력을 받아서 워크플로우를 실행하고 결과를 반환합니다."""
        # messages를 HumanMessage 객체로 변환
        messages = []
        for msg in input_state.get("messages", []):
            if isinstance(msg, dict):
                messages.append(HumanMessage(content=msg.get("content", "")))
            else:
                messages.append(msg)
        
        # 입력 state의 필수 필드들을 확인하고 누락된 필드들을 기본값으로 채움
        complete_state = {
            "messages": messages,
            "doc_type": input_state.get("doc_type"),
            "template_content": input_state.get("template_content"),
            "filled_data": input_state.get("filled_data"),
            "violation": input_state.get("violation"),
            "final_doc": input_state.get("final_doc"),
            "retry_count": input_state.get("retry_count", 0),
            "restart_classification": input_state.get("restart_classification"),
            "classification_retry_count": input_state.get("classification_retry_count"),
            "end_process": input_state.get("end_process"),
            "parse_retry_count": input_state.get("parse_retry_count"),
            "parse_failed": input_state.get("parse_failed")
        }
        
        # 사용자 입력을 메시지에 추가
        complete_state["messages"].append(HumanMessage(content=user_input))
        
        # 그래프 실행
        final_state = self.app.invoke(complete_state)
        
        # 최종 상태가 성공적인지 확인 (규정 위반 없음 + 데이터 존재)
        if final_state.get("violation") == "OK" and final_state.get("filled_data"):
            print("\n" + "="*50)
            print("📄 최종 파싱 결과:")
            print("="*50)
            
            import json
            result = json.dumps(final_state["filled_data"], indent=2, ensure_ascii=False)
            print(result)
            
            return final_state["filled_data"]
        else:
            # 실패 메시지는 각 노드/라우터에서 이미 출력되었으므로, 여기서는 None만 반환하여
            # __main__ 블록에서 "처리 실패"가 출력되도록 함.
            return None

if __name__ == "__main__":
    # 에이전트 실행 예시
    agent = DocumentDraftAgent()

    # 에이전트 실행
    result = agent.run()
    
    if result:
        print("\n✅ 처리 완료!")
        print("반환된 결과:", result)
    else:
        print("\n❌ 처리 실패")