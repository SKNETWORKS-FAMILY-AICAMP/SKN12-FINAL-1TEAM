import os
import re
import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional, TypedDict
from datetime import datetime
from typing_extensions import Annotated
import pandas as pd
import httpx
from dotenv import load_dotenv
from openai import AsyncOpenAI
from langgraph.graph import StateGraph

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from jinja2 import Template


load_dotenv()
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


def OW(_, b):  # OverWrite
    return b

# 3) ReportState에만 적용
class ReportState(TypedDict):
    company_name: str
    start_month: Optional[int]
    end_month: Optional[int]

    final_report: Annotated[Optional[str], OW]
    grade_result: Annotated[Optional[Dict[str, Any]], OW]
    grade_report: Annotated[Optional[str], OW]
    same_grade_report: Annotated[Optional[str], OW]
    growth_report: Annotated[Optional[str], OW]
    strategy_report: Annotated[Optional[str], OW]

    target_df_markdown: Annotated[Optional[str], OW]
    target_df_summary: Annotated[Optional[Dict[str, Any]], OW]


def _to_builtin(o):
        """numpy, pandas 타입을 파이썬 내장형으로 변환"""
        import numpy as np
        if isinstance(o, dict):
            return {k: _to_builtin(v) for k, v in o.items()}
        if isinstance(o, list):
            return [_to_builtin(v) for v in o]
        if isinstance(o, (np.integer,)):
            return int(o)
        if isinstance(o, (np.floating,)):
            return float(o)
        try:
            return o.item()  # pandas/numpy scalar
        except Exception:
            return o

# Agent
# -------------------------------
class ClientAgent:
    #def __init__(self, data_filename: str = "좋은제약_거래처정보.xlsx"):
        #self.data_path = Path(__file__).with_name(data_filename)
        #self.df: pd.DataFrame = pd.DataFrame()
        #self._load_data()

        #api_key = os.getenv("OPENAI_API_KEY")
        #if not api_key:
            #raise ValueError("OPENAI_API_KEY 환경 변수가 설정되지 않았습니다.")
        #self.client = AsyncOpenAI(api_key=api_key)

        # 임계값

        #self.revenue_threshold = {"A": 3000000, "B": 2000000, "C": 1000000, "D": 500000, "E": 100000}
        #self.profit_threshold  = {"A": 10, "B": 15, "C": 20, "D": 25, "E": 30}  # 낮을수록 좋음(reverse=True)
        #self.patience_threshold = {"A": 2200, "B": 1800, "C": 1400, "D": 1000, "E": 500}
        #self.interaction_threshold = {"A": 60, "B": 45, "C": 30, "D": 15, "E": 0}

        # API 설정 (존재하면 API 사용, 없으면 엑셀 사용)
        #self.api_base_url: Optional[str] = os.getenv("API_BASE_URL")
        #self.api_jwt_token: Optional[str] = os.getenv("API_JWT_TOKEN")


    
    #def _load_data(self):
        #try:
            #if not self.data_path.exists():
                #raise FileNotFoundError(f"엑셀 파일을 찾을 수 없습니다: {self.data_path}")

            # ✅ 같은 폴더의 엑셀 로드 (컬럼 이름/공백 정리 포함)
            #df = pd.read_excel(self.data_path, engine="openpyxl")

            # 컬럼명 공백 제거 및 양끝 공백 트림
            #df.columns = [str(c).strip() for c in df.columns]

            #if "월" not in df.columns:
                #raise KeyError("엑셀에 '월' 컬럼이 없습니다.")

            # YYYYMM → datetime → YYYYMM int
            #df["월"] = pd.to_datetime(df["월"].astype(str), format="%Y%m", errors="coerce")
            #df["월_int"] = df["월"].dt.strftime("%Y%m").astype(int)

            #if "거래처ID" not in df.columns:
                #raise KeyError("엑셀에 '거래처ID' 컬럼이 없습니다.")

            #self.df = df
            #logger.info(f"[OK] 데이터 로드: {len(self.df)}건 (경로: {self.data_path})")
        #except Exception as e:
            #logger.error(f"[ERROR] 데이터 로드 실패: {e}")
            #self.df = pd.DataFrame()
            #return
        
    def text2sql_fetch(self, company_name: str,
                       start_month: Optional[int],
                       end_month: Optional[int]) -> pd.DataFrame:
        # API 우선 사용
        if self._is_api_configured():
            try:
                sm = f"{start_month}" if start_month is not None else None
                em = f"{end_month}" if end_month is not None else None
                if not sm or not em:
                    return pd.DataFrame()
                perf = self._fetch_customer_performance(company_name, sm, em)
                df = self._monthly_data_to_dataframe(perf, fallback_customer_id=company_name)
                return df
            except Exception as api_err:
                logger.error(f"[ERROR] API 조회 실패, 엑셀로 폴백: {api_err}")
                # 계속해서 엑셀 폴백 수행

        # 엑셀 폴백
        base = self.df
        if base.empty:
            return pd.DataFrame()
        df = base[base["거래처ID"] == company_name].copy()
        if start_month and end_month:
            df = df[(df["월_int"] >= start_month) & (df["월_int"] <= end_month)]
        return df.sort_values("월_int")

    # ---------------------------
    # API 연동 유틸
    # ---------------------------
    def _is_api_configured(self) -> bool:
        return bool(self.api_base_url and self.api_jwt_token)

    def _api_headers(self) -> Dict[str, str]:
        return {"Authorization": self.api_jwt_token} if self.api_jwt_token else {}

    def _fetch_customer_performance(self, customer_id: str, start_month: str, end_month: str) -> Dict[str, Any]:
        """단일 거래처의 기간 성과 조회. 사양: GET /customer/{id}/performance"""
        if not self.api_base_url:
            raise RuntimeError("API_BASE_URL 미설정")
        url = f"{self.api_base_url.rstrip('/')}/customer/{customer_id}/performance"
        params = {"start_month": start_month, "end_month": end_month}
        with httpx.Client(timeout=30.0) as client:
            resp = client.get(url, params=params, headers=self._api_headers())
            resp.raise_for_status()
            return resp.json()

    def _monthly_data_to_dataframe(self, perf: Dict[str, Any], fallback_customer_id: Optional[str] = None) -> pd.DataFrame:
        """API 응답을 내부 계산용 DataFrame으로 변환하여 반환"""
        monthly = perf.get("monthly_data") or perf.get("data") or []
        df = pd.DataFrame(monthly)
        if df.empty:
            return df

        # 컬럼명 매핑: API '사용예산' → 내부 '사용 예산'
        rename_map = {
            "사용예산": "사용 예산",
            "환자수": "총환자수",
        }
        df = df.rename(columns=rename_map)

        # 필수 컬럼 보강
        for col in ["매출", "사용 예산", "총환자수"]:
            if col not in df.columns:
                df[col] = 0

        # 월 파싱 및 정렬
        month_col = None
        if "월" in df.columns:
            month_col = "월"
        elif "month" in df.columns:
            month_col = "month"

        if month_col:
            df["월"] = pd.to_datetime(df[month_col].astype(str), format="%Y%m", errors="coerce")
            df["월_int"] = df["월"].dt.strftime("%Y%m").astype("Int64")
        else:
            # 월 정보가 없으면 정렬 불가. 그대로 반환
            df["월_int"] = pd.Series([None] * len(df), dtype="Int64")

        # 거래처ID 부여
        if "거래처ID" not in df.columns:
            cid = perf.get("customer_id") or fallback_customer_id
            if cid is not None:
                df["거래처ID"] = cid

        # 정렬 및 정리
        if "월_int" in df.columns:
            df = df.dropna(subset=["월"]) if "월" in df.columns else df
            df = df.sort_values("월_int")
        return df


    # 문서 생성 유틸리티
    # -------------------------------
    def _create_element(self, name):
        """Word 문서에 커스텀 요소를 추가하기 위한 헬퍼 함수"""
        return OxmlElement(name)

    def _add_element_after(self, paragraph, element):
        """문단 다음에 요소를 추가하는 헬퍼 함수"""
        p = paragraph._p
        p.addnext(element)
    
    def _apply_run_font(self, run, size_pt=11, bold=False, italic=False, family='맑은 고딕'):
        run.font.name = family
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic

        r = run._element          
        rPr = r.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.append(rPr)

        rFonts = rPr.rFonts
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)

        rFonts.set(qn('w:eastAsia'), family)
        rFonts.set(qn('w:ascii'), family)
        rFonts.set(qn('w:cs'), family)

    def _add_page_break(self, document):
        """페이지 나누기 추가"""
        document.add_page_break()

    def _add_heading_with_style(self, document, text, level=1):
        """스타일이 적용된 제목 추가"""
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            self._apply_run_font(
                run,
                size_pt=(16 if level == 1 else 14 if level == 2 else 12),
                bold=True,
                italic=False,
                family='맑은 고딕',
            )
        return heading

    def _add_paragraph_with_style(self, document, text, bold=False, italic=False):
        """스타일이 적용된 문단 추가"""
        paragraph = document.add_paragraph(text)
        for run in paragraph.runs:
            self._apply_run_font(
                run,
                size_pt=11,
                bold=bold,
                italic=italic,
                family='맑은 고딕',
            )
        return paragraph

    def _set_default_style_korean(self, document, family='맑은 고딕'):
        """문서 기본 스타일에 한글 폰트(East Asia)를 지정"""
        try:
            normal_style = document.styles['Normal']
            normal_style.font.name = family
            normal_style.font.size = Pt(11)

            rPr = normal_style.element.rPr
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                normal_style.element.append(rPr)

            rFonts = rPr.rFonts
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)

            rFonts.set(qn('w:eastAsia'), family)
            rFonts.set(qn('w:ascii'), family)
            rFonts.set(qn('w:cs'), family)
        except Exception:
            pass

    def _add_table_with_data(self, document, data, headers=None):
        """
        headers: ["열1","열2",...]
        data   : [["a","b",...], ["c","d",...]]
        """
        if headers:
            cols = len(headers)
            rows = 1 + (len(data) if data else 0)
            table = document.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            # 헤더
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = ''
                run = cell.paragraphs[0].add_run(str(h))
                self._apply_run_font(run, size_pt=11, bold=True)
        # 데이터
            if data:
                for i, row_data in enumerate(data, start=1):
                    for j, cell_data in enumerate(row_data):
                        cell = table.rows[i].cells[j]
                        cell.text = ''
                        run = cell.paragraphs[0].add_run(str(cell_data))
                        self._apply_run_font(run, size_pt=11)
        else:
            if not data:
                return None
            cols = len(data[0])
            table = document.add_table(rows=len(data), cols=cols)
            table.style = 'Table Grid'
            for i, row_data in enumerate(data):
                for j, cell_data in enumerate(row_data):
                    cell = table.rows[i].cells[j]
                    cell.text = ''
                    run = cell.paragraphs[0].add_run(str(cell_data))
                    self._apply_run_font(run, size_pt=11)
        return table


    def generate_word_document(self, report_state: ReportState, output_path: Optional[str] = None) -> str:
        """레포트 결과를 Word 문서로 생성"""
        try:
            # 문서 생성 및 기본 스타일 한글 폰트 지정
            doc = Document()
            self._set_default_style_korean(doc, family='맑은 고딕')
            
            # 제목 페이지
            title = doc.add_heading('거래처 분석 보고서', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 부제목
            subtitle = doc.add_paragraph(f"거래처명: {report_state['company_name']}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 생성 날짜
            date_para = doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self._add_page_break(doc)
            
            # 목차
            self._add_heading_with_style(doc, "목차", 1)
            toc_items = [
                "1. 등급 분석 결과",
                "2. 동일 등급 비교 분석", 
                "3. 성장성 분석",
                "4. 영업 전략 제안",
                "5. 종합 분석"
            ]
            for item in toc_items:
                self._add_paragraph_with_style(doc, item)
            
            self._add_page_break(doc)
            
            # 1. 등급 분석 결과
            self._add_heading_with_style(doc, "1. 등급 분석 결과", 1)
            
            if report_state.get('grade_result'):
                grade_result = report_state['grade_result']
                
                # 등급 요약 테이블
                grade_summary = [
                    ["구분", "등급", "점수"],
                    ["매출 등급", grade_result.get('매출등급', 'N/A'), 
                     self.map_grade_to_score(grade_result.get('매출등급', 'E'))],
                    ["수익률 등급", grade_result.get('수익률등급', 'N/A'),
                     self.map_grade_to_score(grade_result.get('수익률등급', 'E'))],
                    ["환자수 등급", grade_result.get('환자수등급', 'N/A'),
                     self.map_grade_to_score(grade_result.get('환자수등급', 'E'))],
                    ["관계도 등급", grade_result.get('관계도등급', 'N/A'),
                     self.map_grade_to_score(grade_result.get('관계도등급', 'E'))],
                    ["최종 등급", grade_result.get('최종등급', 'N/A'), ""]
                ]
                self._add_table_with_data(doc, grade_summary, ["구분", "등급", "점수"])
                
                # 지표 요약
                if '지표요약' in grade_result:
                    self._add_heading_with_style(doc, "주요 지표", 2)
                    indicators = grade_result['지표요약']
                    indicator_data = [
                        ["지표", "값"],
                        ["평균 매출", f"{indicators.get('평균매출', 0):,.0f}원"],
                        ["총 매출", f"{indicators.get('총매출', 0):,.0f}원"],
                        ["총 예산", f"{indicators.get('총예산', 0):,.0f}원"],
                        ["수익률", f"{indicators.get('수익률(%)', 0):.2f}%"],
                        ["평균 환자수", f"{indicators.get('평균환자수', 0):,.0f}명"],
                        ["관계도", f"{indicators.get('관계도(%)', 0):.2f}%"]
                    ]
                    self._add_table_with_data(doc, indicator_data, ["지표", "값"])
            
            # 등급 분석 리포트
            if report_state.get('grade_report'):
                self._add_heading_with_style(doc, "등급 분석 상세", 2)
                self._add_paragraph_with_style(doc, report_state['grade_report'])
            
            self._add_page_break(doc)
            
            # 2. 동일 등급 비교 분석
            self._add_heading_with_style(doc, "2. 동일 등급 비교 분석", 1)
            if report_state.get('same_grade_report'):
                self._add_paragraph_with_style(doc, report_state['same_grade_report'])
            
            self._add_page_break(doc)
            
            # 3. 성장성 분석
            self._add_heading_with_style(doc, "3. 성장성 분석", 1)
            if report_state.get('growth_report'):
                self._add_paragraph_with_style(doc, report_state['growth_report'])
            
            self._add_page_break(doc)
            
            # 4. 영업 전략 제안
            self._add_heading_with_style(doc, "4. 영업 전략 제안", 1)
            if report_state.get('strategy_report'):
                self._add_paragraph_with_style(doc, report_state['strategy_report'])
            
            self._add_page_break(doc)
            
            # 5. 종합 분석
            self._add_heading_with_style(doc, "5. 종합 분석", 1)
            if report_state.get('final_report'):
                self._add_paragraph_with_style(doc, report_state['final_report'])
            
            # 파일 저장
            if output_path is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"거래처분석보고서_{report_state['company_name']}_{timestamp}.docx"
            
            doc.save(output_path)
            logger.info(f"[OK] Word 문서 생성 완료: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"[ERROR] Word 문서 생성 실패: {e}")
            return ""

    def generate_html_document(self, report_state: ReportState, output_path: Optional[str] = None) -> str:
        """레포트 결과를 HTML 문서로 생성 (키/널 안전, Jinja2 사용)"""
        try:
        # 등급 점수 계산(grade_result 없어도 안전)
            gr = report_state.get('grade_result') or {}
            grade_scores = {'매출': 0, '수익률': 0, '환자수': 0, '관계도': 0}
            if gr:
                grade_scores = {
                    '매출':  self.map_grade_to_score(gr.get('매출등급', 'E')),
                    '수익률': self.map_grade_to_score(gr.get('수익률등급', 'E')),
                    '환자수': self.map_grade_to_score(gr.get('환자수등급', 'E')),
                    '관계도': self.map_grade_to_score(gr.get('관계도등급', 'E')),
                }

        # HTML 템플릿 (괄호 포함 키는 대괄호 인덱싱 + .get 사용)
            html_template = """
<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<meta http-equiv="Content-Type" content="text/html; charset=UTF-8" />
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>거래처 분석 보고서 - {{ company_name }}</title>
<style>
body { font-family: 'Malgun Gothic','맑은 고딕',sans-serif; line-height:1.6; margin:0; padding:20px; background:#f5f5f5; }
.container { max-width:1200px; margin:0 auto; background:#fff; padding:40px; box-shadow:0 0 10px rgba(0,0,0,.1); border-radius:8px; }
.header { text-align:center; border-bottom:3px solid #2c3e50; padding-bottom:20px; margin-bottom:30px; }
.header h1 { color:#2c3e50; margin:0; font-size:2.2em; }
.header .subtitle { color:#7f8c8d; font-size:1.1em; margin:10px 0; }
.header .date { color:#95a5a6; font-size:.95em; }
.toc { background:#ecf0f1; padding:20px; border-radius:5px; margin:20px 0; }
.toc h2 { color:#2c3e50; margin:0 0 10px; }
.toc ul { list-style:none; padding-left:0; margin:0; }
.toc li { padding:6px 0; border-bottom:1px solid #bdc3c7; }
.toc li:last-child { border-bottom:none; }
.section { margin:30px 0; page-break-inside:avoid; }
.section h1 { color:#2c3e50; border-left:5px solid #3498db; padding-left:15px; margin-bottom:16px; }
.section h2 { color:#34495e; margin:22px 0 12px; }
table { width:100%; border-collapse:collapse; margin:16px 0; background:#fff; }
th, td { border:1px solid #ddd; padding:10px 12px; text-align:left; }
th { background:#3498db; color:#fff; font-weight:bold; }
tr:nth-child(even) { background:#f7f9fb; }
.grade-a { background:#d5f4e6; }
.grade-b { background:#d4edda; }
.grade-c { background:#fff3cd; }
.grade-d { background:#f8d7da; }
.grade-e { background:#f5c6cb; }
.content { background:#f8f9fa; padding:16px; border-radius:5px; border-left:4px solid #3498db; margin:12px 0; white-space:normal; }
.footer { text-align:center; margin-top:40px; padding-top:18px; border-top:1px solid #ecf0f1; color:#7f8c8d; font-size:.95em; }
@media print { body { background:#fff; } .container { box-shadow:none; } .section { page-break-inside:avoid; } }
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <h1>거래처 분석 보고서</h1>
    <div class="subtitle">거래처명: {{ company_name }}</div>
    <div class="date">생성일: {{ generation_date }}</div>
  </div>

  <div class="toc">
    <h2>목차</h2>
    <ul>
      <li>1. 등급 분석 결과</li>
      <li>2. 동일 등급 비교 분석</li>
      <li>3. 성장성 분석</li>
      <li>4. 영업 전략 제안</li>
      <li>5. 종합 분석</li>
    </ul>
  </div>

  <div class="section">
    <h1>1. 등급 분석 결과</h1>
    {% if grade_result %}
      <h2>등급 요약</h2>
      <table>
        <thead>
          <tr><th>구분</th><th>등급</th><th>점수</th></tr>
        </thead>
        <tbody>
          <tr class="grade-{{ (grade_result.get('매출등급','E')|lower) }}">
            <td>매출 등급</td>
            <td>{{ grade_result.get('매출등급','N/A') }}</td>
            <td>{{ grade_scores.매출 }}</td>
          </tr>
          <tr class="grade-{{ (grade_result.get('수익률등급','E')|lower) }}">
            <td>수익률 등급</td>
            <td>{{ grade_result.get('수익률등급','N/A') }}</td>
            <td>{{ grade_scores.수익률 }}</td>
          </tr>
          <tr class="grade-{{ (grade_result.get('환자수등급','E')|lower) }}">
            <td>환자수 등급</td>
            <td>{{ grade_result.get('환자수등급','N/A') }}</td>
            <td>{{ grade_scores.환자수 }}</td>
          </tr>
          <tr class="grade-{{ (grade_result.get('관계도등급','E')|lower) }}">
            <td>관계도 등급</td>
            <td>{{ grade_result.get('관계도등급','N/A') }}</td>
            <td>{{ grade_scores.관계도 }}</td>
          </tr>
          <tr class="grade-{{ (grade_result.get('최종등급','E')|lower) }}">
            <td><strong>최종 등급</strong></td>
            <td><strong>{{ grade_result.get('최종등급','N/A') }}</strong></td>
            <td></td>
          </tr>
        </tbody>
      </table>

      {% if grade_result.get('지표요약') %}
      <h2>주요 지표</h2>
      <table>
        <thead><tr><th>지표</th><th>값</th></tr></thead>
        <tbody>
          <tr><td>평균 매출</td>
              <td>{{ "{:,.0f}".format(grade_result['지표요약'].get('평균매출', 0)) }}원</td></tr>
          <tr><td>총 매출</td>
              <td>{{ "{:,.0f}".format(grade_result['지표요약'].get('총매출', 0)) }}원</td></tr>
          <tr><td>총 예산</td>
              <td>{{ "{:,.0f}".format(grade_result['지표요약'].get('총예산', 0)) }}원</td></tr>
          <tr><td>수익률</td>
              <td>{{ "{:.2f}".format(grade_result['지표요약'].get('수익률(%)', 0)) }}%</td></tr>
          <tr><td>평균 환자수</td>
              <td>{{ "{:,.0f}".format(grade_result['지표요약'].get('평균환자수', 0)) }}명</td></tr>
          <tr><td>관계도</td>
              <td>{{ "{:.2f}".format(grade_result['지표요약'].get('관계도(%)', 0)) }}%</td></tr>
        </tbody>
      </table>
      {% endif %}
    {% endif %}

    {% if grade_report %}
      <h2>등급 분석 상세</h2>
      <div class="content">{{ grade_report | replace('\n','<br>') | safe }}</div>
    {% endif %}
  </div>

  <div class="section">
    <h1>2. 동일 등급 비교 분석</h1>
    {% if same_grade_report %}
      <div class="content">{{ same_grade_report | replace('\n','<br>') | safe }}</div>
    {% endif %}
  </div>

  <div class="section">
    <h1>3. 성장성 분석</h1>
    {% if growth_report %}
      <div class="content">{{ growth_report | replace('\n','<br>') | safe }}</div>
    {% endif %}
  </div>

  <div class="section">
    <h1>4. 영업 전략 제안</h1>
    {% if strategy_report %}
      <div class="content">{{ strategy_report | replace('\n','<br>') | safe }}</div>
    {% endif %}
  </div>

  <div class="section">
    <h1>5. 종합 분석</h1>
    {% if final_report %}
      <div class="content">{{ final_report | replace('\n','<br>') | safe }}</div>
    {% endif %}
  </div>

  <div class="footer">
    <p>본 보고서는 AI 분석 시스템을 통해 자동 생성되었습니다.</p>
    <p>© 2024 제약사 영업 분석 시스템</p>
  </div>
</div>
</body>
</html>
        """

        # 템플릿 렌더링
            template = Template(html_template)
            html_content = template.render(
                company_name=report_state['company_name'],
                generation_date=datetime.now().strftime('%Y년 %m월 %d일'),
                grade_result=gr if gr else None,
                grade_scores=grade_scores,
                grade_report=report_state.get('grade_report'),
                same_grade_report=report_state.get('same_grade_report'),
                growth_report=report_state.get('growth_report'),
                strategy_report=report_state.get('strategy_report'),
                final_report=report_state.get('final_report')
        )

        # 파일 저장
            if output_path is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"거래처분석보고서_{report_state['company_name']}_{timestamp}.html"

            with open(output_path, 'w', encoding='utf-8-sig') as f:
                f.write(html_content)

            logger.info(f"[OK] HTML 문서 생성 완료: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"[ERROR] HTML 문서 생성 실패: {e}")
            return ""


    def generate_documents(self, report_state: ReportState, output_dir: Optional[str] = None) -> Dict[str, str]:
        """레포트 결과를 Word와 HTML 문서로 모두 생성"""
        try:
            if output_dir is None:
                output_dir = Path.cwd()
            else:
                output_dir = Path(output_dir)
                output_dir.mkdir(parents=True, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            company_name = report_state['company_name']
            
            # Word 문서 생성
            word_filename = f"거래처분석보고서_{company_name}_{timestamp}.docx"
            word_path = output_dir / word_filename
            word_result = self.generate_word_document(report_state, str(word_path))
            
            # HTML 문서 생성
            html_filename = f"거래처분석보고서_{company_name}_{timestamp}.html"
            html_path = output_dir / html_filename
            html_result = self.generate_html_document(report_state, str(html_path))
            
            return {
                "word": word_result,
                "html": html_result,
                "output_dir": str(output_dir)
            }
            
        except Exception as e:
            logger.error(f"[ERROR] 문서 생성 실패: {e}")
            return {"word": "", "html": "", "output_dir": ""}


    # ---------------------------
    # 등급 계산 유틸
    # ---------------------------
    def _get_grade(self, value: float, threshold_dict: Dict[str, float], reverse: bool = False) -> str:
        grades = ["A", "B", "C", "D", "E"]
        if reverse:  # 낮을수록 좋은 지표
            for g in grades:
                if value <= threshold_dict[g]:
                    return g
            return "E"
        else:
            for g in grades:
                if value >= threshold_dict[g]:
                    return g
            return "E"

    def map_grade_to_score(self, grade: str) -> int:
        return {"A": 5, "B": 4, "C": 3, "D": 2, "E": 1}.get(grade.upper(), 0)

    def map_score_to_grade(self, score: int) -> str:
        return {5: "A", 4: "B", 3: "C", 2: "D", 1: "E"}.get(score, "E")

    def calculate_company_grade_from_df(self, df: pd.DataFrame, company_name: str) -> Dict[str, Any]:
        if df.empty:
            return {"error": "데이터 없음", "최종등급": "N/A"}

        avg_revenue = df["매출"].mean() if "매출" in df.columns else 0.0
        total_revenue = df["매출"].sum() if "매출" in df.columns else 0.0
        total_budget = df["사용 예산"].sum() if "사용 예산" in df.columns else 0.0

        profit_rate = (total_budget / total_revenue * 100) if total_revenue > 0 else 0.0

        avg_patients = df["총환자수"].mean() if "총환자수" in df.columns else 0.0
        interaction_rate = ((df["총환자수"].sum() * 30000) / total_revenue * 100) if total_revenue > 0 else 0.0

        revenue_grade = self._get_grade(avg_revenue, self.revenue_threshold, reverse=False)
        profit_grade  = self._get_grade(profit_rate,  self.profit_threshold,  reverse=True)
        patients_grade = self._get_grade(avg_patients, self.patience_threshold, reverse=False)
        interaction_grade = self._get_grade(interaction_rate, self.interaction_threshold, reverse=False)

        scores = {
            "매출액": self.map_grade_to_score(revenue_grade),
            "수익률": self.map_grade_to_score(profit_grade),
            "환자수": self.map_grade_to_score(patients_grade),
            "관계도": self.map_grade_to_score(interaction_grade),
        }
        avg_score = sum(scores.values()) / len(scores)
        final_grade = self.map_score_to_grade(round(avg_score))

        return _to_builtin({
        "거래처명": company_name,
        "매출등급": revenue_grade,
        "수익률등급": profit_grade,
        "환자수등급": patients_grade,
        "관계도등급": interaction_grade,
        "최종등급": final_grade,
        "지표요약": {
        "평균매출": avg_revenue,
        "총매출": total_revenue,
        "총예산": total_budget,
        "수익률(%)": profit_rate,
        "평균환자수": avg_patients,
        "관계도(%)": interaction_rate,
    }
})

    # ---------------------------
    # LLM 공통 호출
    # ---------------------------
    async def use_llm(self, prompt: str) -> str:
        try:
            resp = await self.client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            return (resp.choices[0].message.content or "").strip()
        except Exception as e:
            logger.error(f"LLM 호출 실패: {e}")
            return f"AI 분석 생성 실패: {e}"

    # 2번 노드: 등급계산 + 분석(임계값/공식 설명 포함) 프롬프트 합치기 처음에 fetch함수에서 기간필터링을한 데이터를 가져오느데 여기서 start_month, end_month를 파라미터로 받는이유는 period_str위해
    # ---------------------------
    async def grade_and_analysis(self, company_name: str, target_df: pd.DataFrame,             
                                 start_month: Optional[int], end_month: Optional[int]) -> Dict[str, Any]:
        target_df = self.text2sql_fetch(company_name, start_month, end_month)
        grade_result = self.calculate_company_grade_from_df(target_df, company_name)

        thresholds_pretty = json.dumps({
            "매출 임계값": self.revenue_threshold,
            "수익률 임계값(낮을수록 좋음)": self.profit_threshold,
            "환자수 임계값": self.patience_threshold,
            "관계도 임계값": self.interaction_threshold,
            "사용된 공식": {
                "수익률(%)": "(총예산 / 총매출) * 100",
                "관계도(%)": "(총환자수 합계 * 30000) / 총매출 * 100"
            }
        }, ensure_ascii=False, indent=2)

        pretty = json.dumps(grade_result, ensure_ascii=False, indent=2)
        period_str = f"{start_month or '전체기간'} ~ {end_month or '전체기간'}"

        prompt = f"""
너는 제약사 영업 데이터 분석 전문가다.
아래 **등급 계산 결과**와 **임계값/공식**을 바탕으로, 왜 해당 거래처가 그 등급이 되었는지 분석하라.
분석할 때, 각 지표가 어느 임계값 구간에 해당했는지 근거를 들어 설명하고, 지표 간 트레이드오프도 언급하라.

### 대상
- 거래처명: {company_name}
- 분석 기간: {period_str}

### 등급 계산 결과
{pretty}

### 임계값/공식
{thresholds_pretty}

### 작성 지침
1) 매출/수익률/환자수/관계도 각각에 대해, 실제 지표값이 어느 등급 임계구간에 속하는지 근거를 제시하라.
2) 최종등급 산정 로직(지표 점수 평균→반올림)에 대해 설명하라.
3) 간단한 개선포인트 지정
"""
        grade_report = await self.use_llm(prompt)
        return {"grade_report": grade_report, "grade_result": grade_result}

    # ---------------------------
    # 기존 리포트들 (3/4/5번)
    # ---------------------------
    def get_same_grade_companies(self, company_name: str,
                                 start_month: Optional[int] = None,
                                 end_month: Optional[int] = None) -> Dict[str, Any]:
        # 동일 등급 비교는 전체 df 사용(기간 필터 적용)
        base = self.df.copy()
        if base.empty:
            return {"target_grade": "N/A", "target_info": {}, "same_grade_companies": []}

        target_df = base[base["거래처ID"] == company_name].copy()
        if start_month and end_month:
            base = base[(base["월_int"] >= start_month) & (base["월_int"] <= end_month)]
            target_df = target_df[(target_df["월_int"] >= start_month) & (target_df["월_int"] <= end_month)]

        target_grade_data = self.calculate_company_grade_from_df(target_df, company_name)
        target_grade = target_grade_data.get("최종등급", "N/A")

        same_grade_companies = []
        for cid in base["거래처ID"].dropna().unique():
            if cid == company_name:
                continue
            df_c = base[base["거래처ID"] == cid]
            g = self.calculate_company_grade_from_df(df_c, cid).get("최종등급")
            if g == target_grade:
                dept = str(df_c["과"].iloc[0]) if "과" in df_c.columns and not df_c.empty else "정보없음"
                same_grade_companies.append({
                    "거래처ID": cid,
                    "과": dept,
                    "매출": df_c["매출"].mean() if "매출" in df_c.columns else 0.0,
                    "사용 예산": df_c["사용 예산"].mean() if "사용 예산" in df_c.columns else 0.0
                })

        return {
            "target_grade": target_grade,
            "target_info": target_grade_data,
            "same_grade_companies": same_grade_companies
        }

    def split_companies_by_department(self, company_name: str, same_grade_result: Dict[str, Any]) -> Dict[str, Any]:
        subset = self.df[self.df["거래처ID"] == company_name]
        target_department = str(subset["과"].iloc[0]) if "과" in subset.columns and not subset.empty else "정보없음"

        rows = same_grade_result.get("same_grade_companies", [])
        same_grade_df = pd.DataFrame(rows) if rows else pd.DataFrame(columns=["거래처ID", "과", "매출", "사용 예산"])

        for c in ["과", "매출", "사용 예산"]:
            if c not in same_grade_df.columns:
                same_grade_df[c] = pd.Series(dtype="object" if c == "과" else "float64")

        same_department_df = same_grade_df[same_grade_df["과"] == target_department]
        diff_department_df = same_grade_df[same_grade_df["과"] != target_department]

        return {
            "same_department": same_department_df,
            "diff_department": diff_department_df,
            "target_department": target_department
        }

    async def analyze_same_grade_departments(self, company_name: str,
                                             start_month: Optional[int] = None,
                                             end_month: Optional[int] = None) -> str:
        same_grade_result = self.get_same_grade_companies(company_name, start_month, end_month)
        split = self.split_companies_by_department(company_name, same_grade_result)

        target_df = self.df[self.df["거래처ID"] == company_name].copy()
        if start_month and end_month:
            target_df = target_df[(target_df["월_int"] >= start_month) & (target_df["월_int"] <= end_month)]

        target_avg_revenue = target_df["매출"].mean() if "매출" in target_df.columns else 0.0
        target_avg_budget = target_df["사용 예산"].mean() if "사용 예산" in target_df.columns else 0.0

        same_dept_avg_revenue = split["same_department"]["매출"].mean() if not split["same_department"].empty else 0.0
        diff_dept_avg_budget = split["diff_department"]["사용 예산"].mean() if not split["diff_department"].empty else 0.0

        prompt = f"""
너는 제약사 영업 데이터 분석 전문가이다.
아래 데이터를 기반으로 **대상 병원의 매출과 사용 예산을 같은/다른 진료과 병원들과 비교 분석**하는 보고서를 작성하라.

### 대상 병원
- 이름: {company_name}
- 분석 기간: {start_month or '전체기간'} ~ {end_month or '전체기간'}
- 대상 병원 평균 매출: {target_avg_revenue:,.0f}원
- 대상 병원 평균 사용 예산: {target_avg_budget:,.0f}원
- 진료과: {split['target_department']}

### 비교 그룹
1. 같은 진료과 & 동일 등급 병원 평균 매출: {same_dept_avg_revenue:,.0f}원
2. 다른 진료과 & 동일 등급 병원 평균 사용 예산: {diff_dept_avg_budget:,.0f}원

### 작성 지침
- 우위/열위 지표를 지적하고 개선 포인트를 제안하라.
- 구체적 수치 비교 포함.
"""
        return await self.use_llm(prompt)

    async def generate_growth_analysis_report(self, company_name: str,
                                              start_month: Optional[int] = None,
                                              end_month: Optional[int] = None) -> str:
        company_df = self.df[self.df["거래처ID"] == company_name].copy()
        if start_month and end_month:
            company_df = company_df[(company_df["월_int"] >= start_month) & (company_df["월_int"] <= end_month)]
        if company_df.empty:
            return f"{company_name}의 해당 기간 데이터가 없습니다."

        company_df = company_df.sort_values("월_int")
        cols = [c for c in ["월_int", "매출", "사용 예산", "월방문횟수"] if c in company_df.columns]
        monthly_data = company_df[cols].copy()
        monthly_data_md = monthly_data.to_markdown(index=False)

        start_visits = company_df["월방문횟수"].iloc[0] if "월방문횟수" in company_df.columns else 0.0
        end_visits = company_df["월방문횟수"].iloc[-1] if "월방문횟수" in company_df.columns else 0.0
        avg_visits = company_df["월방문횟수"].mean() if "월방문횟수" in company_df.columns else 0.0

        start_rev = company_df["매출"].iloc[0] if "매출" in company_df.columns else 0.0
        end_rev   = company_df["매출"].iloc[-1] if "매출" in company_df.columns else 0.0
        revenue_growth = (end_rev - start_rev) / start_rev * 100 if start_rev else "데이터 없음"

        start_budget = company_df["사용 예산"].iloc[0] if "사용 예산" in company_df.columns else 0.0
        end_budget   = company_df["사용 예산"].iloc[-1] if "사용 예산" in company_df.columns else 0.0
        budget_growth = ((end_budget - start_budget) / start_budget * 100) if start_budget else "데이터 없음"

        prompt = f"""
너는 제약사 영업 데이터 분석 전문가이며,
아래 월별 데이터를 기반으로 **협력 관계의 성장성**을 분석하라.

### 대상 병원
- 이름: {company_name}
- 분석 기간: {start_month or '전체기간'} ~ {end_month or '전체기간'}

### 월별 데이터 (매출/예산)
{monthly_data_md}

### 참고 지표
- 매출 성장률: {revenue_growth}
- 예산 성장률: {budget_growth}
- 시작 월 방문 횟수: {start_visits}
- 기간 평균 방문 횟수: {avg_visits:.2f}
- 종료 월 방문 횟수: {end_visits}

### 작성 지침
1. 월별 매출 추이와 예산 추이를 함께 분석해, **협력 관계가 균형적으로 성장하고 있는지** 설명하라.
2. 매출만 증가하거나 예산만 증가하는 경우 **관계 불균형 가능성**을 지적하라.
3. 수요(병원 측)와 공급(우리 회사) 관점에서 상호작용 분석을 하라.
4. 방문 횟수 추이를 기반으로 영업 활동 강도 변화(증가/감소/안정)를 평가하라.
"""
        return await self.use_llm(prompt)

    async def generate_sales_strategy_report(
        self,
        company_name: str,
        start_month: Optional[int] = None,
        end_month: Optional[int] = None
    ) -> str:
        df = self.df
        if df is None or df.empty:
            return "데이터가 없습니다."

        # 전체 평균
        overall_avg_patients = float(df["총환자수"].mean(skipna=True)) if "총환자수" in df.columns else 0.0
        overall_avg_visits   = float(df["월방문횟수"].mean(skipna=True)) if "월방문횟수" in df.columns else 0.0
        overall_avg_revenue  = float(df["매출"].mean(skipna=True)) if "매출" in df.columns else 0.0

        # 대상 슬라이스
        company_df = self.text2sql_fetch(company_name, start_month, end_month)
        if company_df.empty:
            return f"{company_name}의 데이터가 존재하지 않습니다."

        company_avg_patients = float(company_df["총환자수"].mean(skipna=True)) if "총환자수" in company_df.columns else 0.0
        company_avg_visits   = float(company_df["월방문횟수"].mean(skipna=True)) if "월방문횟수" in company_df.columns else 0.0
        company_avg_revenue  = float(company_df["매출"].mean(skipna=True)) if "매출" in company_df.columns else 0.0

        # 비율 계산 (직접 연산)
        if overall_avg_patients != 0:
            patients_ratio = ((company_avg_patients - overall_avg_patients) / overall_avg_patients) * 100.0
        else:
            patients_ratio = 0.0

        if overall_avg_revenue != 0:
            revenue_ratio = ((company_avg_revenue - overall_avg_revenue) / overall_avg_revenue) * 100.0
        else:
            revenue_ratio = 0.0

        if overall_avg_visits != 0:
            visits_ratio = ((company_avg_visits - overall_avg_visits) / overall_avg_visits) * 100.0
        else:
            visits_ratio = 0.0

        # 최근 3개월 대비 자체 평균
        if company_avg_revenue != 0:
            recent_revenue = float(company_df["매출"].tail(3).mean(skipna=True)) if "매출" in company_df.columns else 0.0
            recent_ratio = ((recent_revenue - company_avg_revenue) / company_avg_revenue) * 100.0
        else:
            recent_revenue = 0.0
            recent_ratio = 0.0

        # 인사이트 도출
        insights = []
        if patients_ratio > 15 and revenue_ratio < 0:
            insights.append("환자 수는 높지만 매출은 평균 이하로, 미개척 잠재 시장 가능성이 있습니다.")
        if patients_ratio > 15 and revenue_ratio > 15:
            insights.append("환자 수와 매출 모두 평균 이상으로 핵심 고객군에 속합니다.")
        if visits_ratio > 15 and revenue_ratio < 0:
            insights.append("방문 횟수는 많으나 매출이 낮아 방문 전략 재검토가 필요합니다.")
        if recent_ratio > 15:
            insights.append("최근 3개월 매출이 급상승, 성장세 유지 전략 필요.")
        if recent_ratio < -15:
            insights.append("최근 3개월 매출이 급감, 원인 파악 및 리커버리 전략 필요.")
        if not insights:
            insights.append("모든 지표가 평균 범위 내에 있어 안정적이나, 추가 성장 기회 탐색 필요.")
        
        prompt = f"""
너는 제약사 영업전략 컨설턴트                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                 이며, 
아래 데이터를 기반으로 **해당 거래처의 영업전략 분석 보고서**를 작성하라.                  

### 대상 거래처                                                                                                     
- 이름: {company_name}                                                                                 
- 분석 기간: 전체 기간
                                                                                        
### 전체 거래처 평균
- 환자수 평균: {overall_avg_patients:.1f}
- 방문 횟수 평균: {overall_avg_visits:.1f}
- 매출 평균: {overall_avg_revenue:.1f}

### 해당 거래처 평균
- 환자수 평균: {company_avg_patients:.1f}                     
- 방문 횟수 평균: {company_avg_visits:.1f}
- 매출 평균: {company_avg_revenue:.1f}                                                              


### 분석 포인트
- {', '.join(insights) if insights else '특이사항 없음'}

### 작성 지침
1. 제공된 insights 목록을 가장 핵심 분석 포인트로 삼아, 각각의 의미를 해석하고 영향 요인을 설명하라.
2. insights에서 드러난 문제점·기회 요인을 기반으로 원인과 배경을 구체적으로 추론하라.
3. insights에 따라 실행 가능한 영업 전략(방문 전략, 제품 포트폴리오, 예산 배분, 프로모션 등)을 제안하라.
4. 숫자 지표(환자수, 방문횟수, 매출)와 insights 해석을 연결해 논리적으로 서술하라.
5. 만약 insights가 '모든 지표가 평균 범위 내에 있어 안정적이나, 추가 성장 기회 탐색 필요.'일 경우,
    전반적으로 안정적인 고객으로 평가하고 관계 유지 및 성장 가능성 탐색 전략에 중점 두어 작성하라.
"""
        return await self.use_llm(prompt)


# Graph (Start → 2/3/4/5 동시 → Merge)

def build_full_graph(agent: ClientAgent):
    graph = StateGraph(ReportState)

    # ✅ Start: 1단계(Text2SQL) — 데이터 확보 & 상태 저장
    async def start_node(state: ReportState):
        company = state["company_name"]
        s, e = state["start_month"], state["end_month"]

        target_df = agent.text2sql_fetch(company, s, e)
        target_df_md = (
            target_df[
                ["월_int", "매출", "사용 예산", "총환자수", "월방문횟수"]
                if set(["월_int","매출","사용 예산","총환자수","월방문횟수"]).issubset(target_df.columns)
                else target_df.columns
            ].to_markdown(index=False)
            if not target_df.empty else "데이터 없음"
        )

        state["target_df_markdown"] = target_df_md   
        return state

    # ✅ 2번: 등급 계산+분석 (합치기)
    async def grade_and_analysis_node(state: ReportState):
        company = state["company_name"]
        s, e = state["start_month"], state["end_month"]
        target_df = agent.text2sql_fetch(company, s, e)  # Start에서 가져왔지만 안전하게 재조회
        out = await agent.grade_and_analysis(company, target_df, s, e)
        state["grade_result"] = out["grade_result"]
        state["grade_report"] = out["grade_report"]
        return state

    # 3/4/5번: 동일 등급 / 성장성 / 전략
    async def same_grade_node(state: ReportState):
        company = state["company_name"]
        s, e = state["start_month"], state["end_month"]
        report = await agent.analyze_same_grade_departments(company, s, e)
        return {"same_grade_report": report}

    async def growth_node(state: ReportState):
        company = state["company_name"]
        s, e = state["start_month"], state["end_month"]
        report = await agent.generate_growth_analysis_report(company, s, e)
        return {"growth_report": report}

    async def strategy_node(state: ReportState):
        company = state["company_name"]
        report = await agent.generate_sales_strategy_report(company)
        return {"strategy_report": report}

    # Merge
    async def merge_node(state: ReportState):
        grade_result = json.dumps(state.get("grade_result", {}), ensure_ascii=False, indent=2)
        prompt = f"""
너는 제약사 영업 데이터 컨설턴트이다.
아래 2~5번 리포트를 종합하여 통합 보고서를 작성하라.

### (2) 등급 분석
**등급 계산 결과**
{grade_result}
{state.get('grade_report')}

### (3) 동일 등급 비교 분석
{state.get('same_grade_report')}

### (4) 성장성 분석
{state.get('growth_report')}

### (5) 영업 전략 제안
{state.get('strategy_report')}
"""
        merged = await agent.use_llm(prompt)
        state["final_report"] = merged
        return state

    
    graph.add_node("start", start_node)
    graph.add_node("grade_and_analysis", grade_and_analysis_node)  
    graph.add_node("same_grade", same_grade_node)                  
    graph.add_node("growth", growth_node)                          
    graph.add_node("strategy", strategy_node)                      
    graph.add_node("merge", merge_node)

    
    graph.set_entry_point("start")
    graph.add_edge("start", "grade_and_analysis")
    graph.add_edge("start", "same_grade")
    graph.add_edge("start", "growth")
    graph.add_edge("start", "strategy")

    graph.add_edge("grade_and_analysis", "merge")
    graph.add_edge("same_grade", "merge")
    graph.add_edge("growth", "merge")
    graph.add_edge("strategy", "merge")

    graph.set_finish_point("merge")
    return graph.compile()

# -------------------------------
# 파이프라인 엔트리
# -------------------------------
async def run_full_pipeline(agent: ClientAgent,
                            company_name: str,
                            start_month: Optional[int] = None,
                            end_month: Optional[int] = None,
                            generate_docs: bool = True,
                            output_dir: Optional[str] = None) -> Dict[str, Any]:
    initial: ReportState = {
        "company_name": company_name,
        "start_month": start_month,
        "end_month": end_month,
        "final_report": None,
        "grade_result": None,
        "grade_report": None,
        "same_grade_report": None,
        "growth_report": None,
        "strategy_report": None,    
        "target_df_markdown": None,
        "target_df_summary": None,
    }
    graph = build_full_graph(agent)
    result_state = await graph.ainvoke(initial)
    
    # 문서 생성
    doc_results = {}
    if generate_docs:
        doc_results = agent.generate_documents(result_state, output_dir)
    
    return {
        "report_state": result_state,
        "documents": doc_results
    }


if __name__ == "__main__":
    import asyncio

    # 분석할 거래처 ID와 기간 지정
    COMPANY_NAME = "미라클신경과의원(강서구 화곡동)"
    START_MONTH = 202401            # YYYYMM 형식, 없으면 None
    END_MONTH = 202403              # YYYYMM 형식, 없으면 None

    # 에이전트 생성
    agent = ClientAgent()

    # 파이프라인 실행
    result = asyncio.run(
        run_full_pipeline(
            agent,
            company_name=COMPANY_NAME,
            start_month=START_MONTH,
            end_month=END_MONTH,
            generate_docs=True,     # Word, HTML 보고서 생성 여부
            output_dir="./output"   # 결과 저장 폴더
        )
    )

    # 결과 출력
    print("\n=== 분석 결과 요약 ===")
    print(result["report_state"]["final_report"])
    print("\n문서 저장 위치:", result["documents"])
